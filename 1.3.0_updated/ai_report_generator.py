# # from __future__ import annotations

# # import csv
# # import math
# # import re
# # import textwrap
# # from pathlib import Path
# # from typing import Any, Callable, Dict, List, Optional, Tuple

# # import pandas as pd
# # import requests
# # from docx import Document
# # from docx.enum.section import WD_ORIENT
# # from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
# # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # from docx.oxml import OxmlElement
# # from docx.oxml.ns import qn
# # from docx.shared import Inches, Pt

# # from asset_registry import AssetRegistry
# # from metadata_utils import PARAM_PRETTY, canonical_metric_key, clean_router_name, normalize_band_value, normalize_floor_name

# # _PARAM_UNITS = {
# #     "signal_strength": "dBm",
# #     "secondary_signal_strength": "dBm",
# #     "tertiary_signal_strength": "dBm",
# #     "snr": "dB",
# #     "noise": "dBm",
# #     "data_rate": "Mbps",
# #     "throughput": "Mbps",
# #     "channel_utilization": "%",
# #     "channel_interference": "dB",
# #     "channel_width": "MHz",
# #     "spectrum_channel_power": "dBm",
# #     "network_health": "score",
# #     "network_issues": "count",
# #     "number_of_access_points": "count",
# # }

# # _HIGHER_IS_BETTER = {
# #     "signal_strength": True,
# #     "secondary_signal_strength": True,
# #     "tertiary_signal_strength": True,
# #     "snr": True,
# #     "noise": False,
# #     "data_rate": True,
# #     "throughput": True,
# #     "channel_utilization": False,
# #     "channel_interference": False,
# #     "channel_width": True,
# #     "spectrum_channel_power": True,
# #     "network_health": True,
# #     "network_issues": False,
# #     "number_of_access_points": False,
# # }


# # def _safe_slug(text: str) -> str:
# #     return re.sub(r"[^A-Za-z0-9_-]+", "_", str(text)).strip("_")


# # def _get_param_info(param_key: str) -> Tuple[str, str, bool]:
# #     pkey = canonical_metric_key(param_key) or param_key
# #     return (
# #         PARAM_PRETTY.get(pkey, pkey.replace("_", " ").title()),
# #         _PARAM_UNITS.get(pkey, ""),
# #         _HIGHER_IS_BETTER.get(pkey, True),
# #     )


# # def call_ollama(prompt: str, model: str = "gemma3:4b", base_url: str = "http://localhost:11434", timeout: int = 90) -> str:
# #     errors = []

# #     # 1) Ollama legacy endpoint (/api/generate)
# #     try:
# #         resp = requests.post(
# #             f"{base_url.rstrip('/')}/api/generate",
# #             json={
# #                 "model": model,
# #                 "prompt": prompt,
# #                 "stream": False,
# #                 "options": {"temperature": 0.2, "top_p": 0.9, "num_predict": 700},
# #             },
# #             timeout=timeout,
# #         )
# #         if resp.status_code == 404:
# #             raise requests.exceptions.HTTPError(f"404 at {base_url.rstrip('/')}/api/generate")
# #         resp.raise_for_status()
# #         body = resp.json()
# #         if isinstance(body, dict) and "response" in body and body.get("response"):
# #             return _strip_preamble(str(body.get("response", "")).strip())
# #     except Exception as exc:
# #         errors.append(f"legacy /api/generate: {exc}")

# #     # 2) Ollama v1 chat completion endpoint (/v1/chat/completions)
# #     try:
# #         resp = requests.post(
# #             f"{base_url.rstrip('/')}/v1/chat/completions",
# #             json={
# #                 "model": model,
# #                 "messages": [{"role": "user", "content": prompt}],
# #                 "temperature": 0.2,
# #                 "max_tokens": 700,
# #             },
# #             timeout=timeout,
# #         )
# #         resp.raise_for_status()
# #         body = resp.json()
# #         choices = body.get("choices", [])
# #         if choices:
# #             content = choices[0].get("message", {}).get("content")
# #             if content:
# #                 return _strip_preamble(str(content).strip())
# #     except Exception as exc:
# #         errors.append(f"v1/chat/completions: {exc}")

# #     # 3) OpenAI-compatible completions endpoint (/v1/completions)
# #     try:
# #         resp = requests.post(
# #             f"{base_url.rstrip('/')}/v1/completions",
# #             json={
# #                 "model": model,
# #                 "prompt": prompt,
# #                 "temperature": 0.2,
# #                 "max_tokens": 700,
# #             },
# #             timeout=timeout,
# #         )
# #         resp.raise_for_status()
# #         body = resp.json()
# #         choices = body.get("choices", [])
# #         if choices:
# #             text = choices[0].get("text") or choices[0].get("message", {}).get("content")
# #             if text:
# #                 return _strip_preamble(str(text).strip())
# #     except Exception as exc:
# #         errors.append(f"v1/completions: {exc}")

# #     return f"[AI: {' | '.join(errors)}]"


# # def _strip_preamble(text: str) -> str:
# #     """Remove common AI preambles and conclusions from response."""
# #     if not text:
# #         return text
# #     text = text.strip()
# #     # Remove leading preamble phrases
# #     preamble_patterns = [
# #         r"^Here'?s?\s+(?:a\s+)?(?:comparative\s+)?(?:report|summary|analysis)[^:]*:\s*",
# #         r"^Based on\s+the\s+provided\s+data[^:]*:\s*",
# #         r"^Here are\s+",
# #         r"^Here'?s?\s+",
# #     ]
# #     for pattern in preamble_patterns:
# #         text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
# #     text = text.strip()
    
# #     # Remove trailing conclusion phrases
# #     conclusion_patterns = [
# #         r"\n*(?:This|These|The above)\s+(?:analysis|summary|points?|bullets?).*$",
# #         r"\n*(?:Overall|In summary|In conclusion)[^.]*$",
# #         r"\n*Feel free to.*$",
# #     ]
# #     for pattern in conclusion_patterns:
# #         text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
# #     return text.strip()


# # def _parse_bullets(text: str) -> List[str]:
# #     text = _strip_preamble(text)
# #     lines: List[str] = []
# #     for line in (text or "").splitlines():
# #         line = line.strip()
# #         if not line:
# #             continue
# #         # Remove leading bullet markers
# #         line = re.sub(r"^[\-•*\d\.\)\s]+", "", line)
# #         # Remove markdown bold/italic markers
# #         line = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", line)
# #         line = line.strip()
# #         if len(line) >= 10:
# #             lines.append(line)
# #     return lines


# # def _fallback_bullets(stats_rows: List[Dict[str, Any]], param_pretty: str, param_unit: str, floor_name: str, band: str, higher_is_better: bool) -> List[str]:
# #     if not stats_rows:
# #         return [f"No usable data was available for {floor_name} on {band}."]

# #     best = stats_rows[0]
# #     worst = stats_rows[-1]
# #     spread = best["avg"] - worst["avg"]
# #     ranking = " > ".join(r["router"] for r in stats_rows)
# #     trend = "more stable" if max(r["range"] for r in stats_rows) < max(abs(best["avg"]) * 0.15, 1) else "more variable"
# #     cluster = None
# #     if len(stats_rows) >= 3:
# #         mids = stats_rows[1:-1]
# #         cluster = ", ".join(f"{r['router']} ({r['avg']:.1f} {param_unit})" for r in mids)

# #     direction = "higher" if higher_is_better else "lower"
# #     impairment = "attenuation, interference, or backhaul efficiency" if str(band).startswith("5") else "radio design, sensitivity, or congestion handling"
# #     bullets = [
# #         f"Across {floor_name} on {band}, the curves are {trend} than in the strongest cases, so router design matters more than distance alone.",
# #         f"{best['router']} delivers the best average {param_pretty} at {best['avg']:.1f} {param_unit}.",
# #     ]
# #     if cluster:
# #         bullets.append(f"The middle tier is formed by {cluster}.")
# #     bullets.extend(
# #         [
# #             f"{worst['router']} is the weakest result at {worst['avg']:.1f} {param_unit}, creating a best-to-worst gap of {spread:.1f} {param_unit}.",
# #             f"The separation likely reflects differences in {impairment}, especially where {direction} {param_pretty.lower()} is preferred.",
# #             f"Overall ranking: {ranking}.",
# #         ]
# #     )
# #     return bullets


# # def _analysis_prompt(param_pretty: str, param_unit: str, floor_name: str, band: str, config_label: str, stats_rows: List[Dict[str, Any]], higher_is_better: bool) -> str:
# #     direction_word = "higher is better" if higher_is_better else "lower is better"
# #     rows_text = "\n".join(
# #         f"{i+1}. {r['router']}: avg={r['avg']:.1f} {param_unit}, min={r['min']:.1f}, max={r['max']:.1f}, spread={r['range']:.1f}"
# #         for i, r in enumerate(stats_rows)
# #     )
# #     return textwrap.dedent(
# #         f"""
# #         You are writing a professional Wi-Fi comparative report. Do NOT add preamble or conclusions.
# #         Write ONLY exactly 5 bullet points. Each bullet must start with •

# #         Metric: {param_pretty} ({param_unit})
# #         Configuration: {config_label}
# #         Floor: {floor_name}
# #         Band: {band}
# #         Interpretation: {direction_word}

# #         Router statistics, best to worst:
# #         {rows_text}

# #         Rules:
# #         - Use only the provided statistics. Do not infer, estimate, or invent rankings, causes, or values.
# #         - Output exactly 5 bullets, and nothing else.
# #         - Every bullet must begin with •
# #         - Bullet 1: summarize the overall pattern and degree of variability across all routers.
# #         - Bullet 2: identify the best performer as the first router in the list, and report its exact max and average value.
# #         - Bullet 3: identify the mid-tier router or routers only if there are 3 or more total routers; include each name and exact max and  average value.
# #         - Bullet 4: identify the worst performer as the last router in the list, and report its exact max and  average value plus the exact gap from the best performer.
# #         - Bullet 5: explain the likely technical basis for the separation using radio-frequency or network-performance principles, and keep the explanation scientifically grounded.
# #         - Do not add any preamble, closing summary, headings, or extra commentary outside the 5 bullets.
# #         - Preserve router names and numeric values exactly as provided.
# #         """
# #     ).strip()


# # def discover_rvr_assets(root: Path, metric_folder: str) -> Dict[str, Any]:
# #     metric_dir = root / metric_folder
# #     plots_actual: Dict[Tuple[str, str], Path] = {}
# #     table_path = None
# #     tables_dir = metric_dir / "tables"
# #     if tables_dir.exists():
# #         candidates = sorted(tables_dir.glob("*_curve_tables.csv"))
# #         if candidates:
# #             table_path = candidates[0]
# #     plot_dir = metric_dir / "plots_actual" / "COMPARE_ROUTERS_BY_FLOOR"
# #     if plot_dir.exists():
# #         for png in sorted(plot_dir.glob("*.png")):
# #             m = re.match(r"^.+?__act__(?P<band>.+?)__floor_(?P<floor>.+?)\.png$", png.name)
# #             if m:
# #                 plots_actual[(m.group("band"), m.group("floor"))] = png
# #     return {"mode": "rvr", "metric_dir": metric_dir, "curve_table_path": table_path, "plots_actual": plots_actual}


# # def discover_mesh_assets(root: Path, metric_folder: str) -> Dict[str, Any]:
# #     metric_dir = root / metric_folder
# #     plot_map: Dict[Tuple[str, str, str], Path] = {}
# #     table_path = None
# #     tables_dir = metric_dir / "tables"
# #     if tables_dir.exists():
# #         candidates = sorted(tables_dir.glob("*_mesh_curve_tables.csv"))
# #         if candidates:
# #             table_path = candidates[0]
# #     for png in sorted(metric_dir.rglob("*.png")):
# #         if png.parent == metric_dir:
# #             continue
# #         router = clean_router_name(png.parent.parent.name) if png.parent.parent != metric_dir else ""
# #         floor_name = normalize_floor_name(png.parent.name)
# #         stem = png.stem
# #         if "_" not in stem:
# #             continue
# #         band, _ = stem.split("_", 1)
# #         plot_map[(router, floor_name, normalize_band_value(band))] = png
# #     return {"mode": "mesh_compare", "metric_dir": metric_dir, "curve_table_path": table_path, "plots_actual": plot_map}


# # def _find_rvr_plot(plot_map: Dict[Tuple[str, str], Path], band: str, floor_name: str) -> Optional[Path]:
# #     band = normalize_band_value(band)
# #     floor_norm = _safe_slug(normalize_floor_name(floor_name))
# #     key = (band, floor_norm)
# #     if key in plot_map:
# #         return plot_map[key]
# #     for (bs, fs), path in plot_map.items():
# #         if band in bs and floor_norm.lower() in fs.lower():
# #             return path
# #     return None


# # def _find_mesh_plot(plot_map: Dict[Tuple[str, str, str], Path], router_key: str, floor_name: str, band: str) -> Optional[Path]:
# #     key = (clean_router_name(router_key), normalize_floor_name(floor_name), normalize_band_value(band))
# #     if key in plot_map:
# #         return plot_map[key]
# #     return None


# # def _compute_stats(subset: pd.DataFrame, higher_is_better: bool) -> List[Dict[str, Any]]:
# #     y_col = "p50" if "p50" in subset.columns else ("mean" if "mean" in subset.columns else None)
# #     if not y_col:
# #         return []
# #     rows: List[Dict[str, Any]] = []
# #     for router_key in subset["router_key"].astype(str).unique():
# #         part = subset[subset["router_key"].astype(str) == router_key]
# #         s = pd.to_numeric(part[y_col], errors="coerce").dropna()
# #         if s.empty:
# #             continue
# #         display = clean_router_name(str(part.get("router_display", pd.Series([router_key])).iloc[0]))
# #         rows.append(
# #             {
# #                 "router": display,
# #                 "router_key": clean_router_name(router_key),
# #                 "avg": float(s.mean()),
# #                 "min": float(s.min()),
# #                 "max": float(s.max()),
# #                 "range": float(s.max() - s.min()),
# #             }
# #         )
# #     rows.sort(key=lambda r: r["avg"], reverse=higher_is_better)
# #     return rows


# # def _set_landscape(section) -> None:
# #     section.orientation = WD_ORIENT.LANDSCAPE
# #     section.page_width, section.page_height = section.page_height, section.page_width
# #     section.left_margin = Inches(0.6)
# #     section.right_margin = Inches(0.6)
# #     section.top_margin = Inches(0.7)
# #     section.bottom_margin = Inches(0.7)


# # def _set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
# #     tc = cell._tc
# #     tcPr = tc.get_or_add_tcPr()
# #     tcMar = tcPr.first_child_found_in("w:tcMar")
# #     if tcMar is None:
# #         tcMar = OxmlElement("w:tcMar")
# #         tcPr.append(tcMar)
# #     for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
# #         element = tcMar.find(qn(f"w:{tag}"))
# #         if element is None:
# #             element = OxmlElement(f"w:{tag}")
# #             tcMar.append(element)
# #         element.set(qn("w:w"), str(value))
# #         element.set(qn("w:type"), "dxa")


# # def _add_heading(doc: Document, text: str, level: int = 1) -> None:
# #     p = doc.add_heading(level=level)
# #     run = p.add_run(text)
# #     run.bold = True


# # def _add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
# #     p = doc.add_paragraph()
# #     run = p.add_run(text)
# #     run.italic = italic
# #     run.font.size = Pt(10.5)


# # def _add_plot(doc: Document, plot_path: Path) -> None:
# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     p.add_run().add_picture(str(plot_path), width=Inches(9.2))


# # def _add_router_cards(doc: Document, router_cards: List[Dict[str, Any]]) -> None:
# #     if not router_cards:
# #         return
# #     cols = 2
# #     rows = math.ceil(len(router_cards) / cols)
# #     table = doc.add_table(rows=rows, cols=cols)
# #     table.alignment = WD_TABLE_ALIGNMENT.CENTER
# #     table.style = "Table Grid"
# #     table.autofit = False
# #     card_width = 4.85
# #     heatmap_width = 4.35
# #     scale_width = 4.0

# #     for idx, card in enumerate(router_cards):
# #         cell = table.cell(idx // cols, idx % cols)
# #         cell.width = Inches(card_width)
# #         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
# #         _set_cell_margins(cell, top=90, start=90, bottom=90, end=90)

# #         p = cell.paragraphs[0]
# #         p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #         name_run = p.add_run(card["display"])
# #         name_run.bold = True
# #         name_run.font.size = Pt(10)

# #         if card.get("scenario_label"):
# #             label_p = cell.add_paragraph()
# #             label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #             label_run = label_p.add_run(card["scenario_label"])
# #             label_run.italic = True
# #             label_run.font.size = Pt(9)

# #         if card.get("heatmap") and Path(card["heatmap"]).exists():
# #             hp = cell.add_paragraph()
# #             hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #             hp.add_run().add_picture(str(card["heatmap"]), width=Inches(heatmap_width))
# #         else:
# #             miss = cell.add_paragraph()
# #             miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #             miss.add_run("Exact heatmap match not found").italic = True

# #         if card.get("scale") and Path(card["scale"]).exists():
# #             sp = cell.add_paragraph()
# #             sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #             sp.add_run().add_picture(str(card["scale"]), width=Inches(scale_width))
# #         else:
# #             miss = cell.add_paragraph()
# #             miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #             miss.add_run("Matching color scale not found").italic = True

# #     total_cells = rows * cols
# #     for idx in range(len(router_cards), total_cells):
# #         table.cell(idx // cols, idx % cols).text = ""


# # def _add_bullets(doc: Document, bullets: List[str]) -> None:
# #     for bullet in bullets:
# #         p = doc.add_paragraph(style="List Bullet")
# #         run = p.add_run(bullet)
# #         run.font.size = Pt(10.5)


# # def _collect_rvr_router_cards(registry: AssetRegistry, subset: pd.DataFrame, parameter_key: str, floor_name: str, band: str) -> List[Dict[str, Any]]:
# #     cards: List[Dict[str, Any]] = []
# #     for router_key in subset["router_key"].astype(str).unique():
# #         part = subset[subset["router_key"].astype(str) == router_key]
# #         display = clean_router_name(str(part["router_display"].iloc[0] if "router_display" in part.columns else router_key))
# #         heatmap, scale = registry.get_pair(router_key, parameter_key, floor_name, band)
# #         cards.append({"router_key": clean_router_name(router_key), "display": display, "heatmap": heatmap, "scale": scale})
# #     return cards


# # def _collect_mesh_cards(with_registry: AssetRegistry, without_registry: AssetRegistry, router_key: str, parameter_key: str, floor_name: str, band: str) -> List[Dict[str, Any]]:
# #     display = clean_router_name(router_key)
# #     with_heatmap, with_scale = with_registry.get_pair(router_key, parameter_key, floor_name, band)
# #     wo_heatmap, wo_scale = without_registry.get_pair(router_key, parameter_key, floor_name, band)
# #     return [
# #         {"router_key": display, "display": display, "scenario_label": "With mesh", "heatmap": with_heatmap, "scale": with_scale},
# #         {"router_key": display, "display": display, "scenario_label": "Without mesh", "heatmap": wo_heatmap, "scale": wo_scale},
# #     ]


# # def _append_audit_rows(audit_rows: List[Dict[str, Any]], cards: List[Dict[str, Any]], parameter_key: str, floor_name: str, band: str, section_type: str):
# #     for card in cards:
# #         heatmap = card.get("heatmap")
# #         scale = card.get("scale")
# #         status = "exact_pair"
# #         if not heatmap and not scale:
# #             status = "missing_heatmap_and_scale"
# #         elif not heatmap:
# #             status = "missing_heatmap"
# #         elif not scale:
# #             status = "missing_scale"
# #         audit_rows.append(
# #             {
# #                 "section_type": section_type,
# #                 "router_key": card.get("router_key", ""),
# #                 "display": card.get("display", ""),
# #                 "scenario_label": card.get("scenario_label", ""),
# #                 "parameter_key": canonical_metric_key(parameter_key) or parameter_key,
# #                 "floor_name": normalize_floor_name(floor_name),
# #                 "band": normalize_band_value(band),
# #                 "heatmap_path": str(heatmap) if heatmap else "",
# #                 "scale_path": str(scale) if scale else "",
# #                 "status": status,
# #             }
# #         )


# # def _write_asset_audit(output_path: Path, audit_rows: List[Dict[str, Any]]) -> Optional[Path]:
# #     if not audit_rows:
# #         return None
# #     audit_path = output_path.with_name(f"{output_path.stem}_asset_audit.csv")
# #     with audit_path.open("w", encoding="utf-8", newline="") as handle:
# #         writer = csv.DictWriter(handle, fieldnames=list(audit_rows[0].keys()))
# #         writer.writeheader()
# #         writer.writerows(audit_rows)
# #     return audit_path


# # def generate_report(
# #     rvr_outputs_root: Path,
# #     extracted_root: Path,
# #     output_path: Path,
# #     metric_folders: List[str],
# #     config_label: str = "Standard",
# #     ai_model: str = "gemma3:4b",
# #     ai_base_url: str = "http://localhost:11434",
# #     use_ai: bool = True,
# #     progress_cb: Optional[Callable[[int, int, str], None]] = None,
# #     mode: str = "rvr",
# #     compare_outputs_root: Optional[Path] = None,
# #     extracted_roots_by_scenario: Optional[Dict[str, Path]] = None,
# #     csv_outputs_root: Optional[Path] = None,
# # ) -> Path:
# #     doc = Document()
# #     _set_landscape(doc.sections[0])
# #     title = doc.add_paragraph()
# #     title.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     run = title.add_run("Wi-Fi Comparative Analysis Report")
# #     run.bold = True
# #     run.font.size = Pt(18)
# #     subtitle = "Mesh vs No Mesh" if mode == "mesh_compare" else "Router-to-router comparison"
# #     _add_paragraph(doc, subtitle, italic=True)

# #     total = max(1, len(metric_folders))
# #     registry = AssetRegistry.from_roots(extracted_root=extracted_root, csv_outputs_root=csv_outputs_root)
# #     with_registry = without_registry = AssetRegistry([])
# #     if mode == "mesh_compare" and extracted_roots_by_scenario:
# #         with_registry = AssetRegistry.from_roots(extracted_root=extracted_roots_by_scenario.get("with_mesh", Path(".")))
# #         without_registry = AssetRegistry.from_roots(extracted_root=extracted_roots_by_scenario.get("without_mesh", Path(".")))

# #     audit_rows: List[Dict[str, Any]] = []

# #     for idx, metric_folder in enumerate(metric_folders, start=1):
# #         metric_folder = canonical_metric_key(metric_folder) or metric_folder
# #         if progress_cb:
# #             progress_cb(idx, total, f"Preparing {metric_folder}")
# #         param_pretty, param_unit, higher_is_better = _get_param_info(metric_folder)
# #         heading_suffix = f" — {config_label}" if config_label and config_label.strip().lower() not in {"", "standard"} else ""
# #         _add_heading(doc, f"Range vs {param_pretty}{heading_suffix}", level=1)

# #         if mode == "mesh_compare":
# #             if compare_outputs_root is None:
# #                 raise ValueError("compare_outputs_root is required for mesh_compare mode")
# #             assets = discover_mesh_assets(compare_outputs_root, metric_folder)
# #         else:
# #             assets = discover_rvr_assets(rvr_outputs_root, metric_folder)

# #         curve_table_path = assets.get("curve_table_path")
# #         if not curve_table_path or not Path(curve_table_path).exists():
# #             _add_paragraph(doc, f"No curve table was found for {param_pretty}.")
# #             continue

# #         curve_table = pd.read_csv(curve_table_path)
# #         if curve_table.empty:
# #             _add_paragraph(doc, f"The curve table for {param_pretty} is empty.")
# #             continue
# #         if "band" in curve_table.columns:
# #             curve_table["band"] = curve_table["band"].astype(str).map(normalize_band_value)
# #         if "floor_name" in curve_table.columns:
# #             curve_table["floor_name"] = curve_table["floor_name"].astype(str).map(normalize_floor_name)
# #         if "router_key" in curve_table.columns:
# #             curve_table["router_key"] = curve_table["router_key"].astype(str).map(clean_router_name)
# #         if "router_display" in curve_table.columns:
# #             curve_table["router_display"] = curve_table["router_display"].astype(str).map(clean_router_name)

# #         intro = f"This section compares {param_pretty} across the available floors and bands." if mode == "rvr" else f"This section compares with-mesh and without-mesh behavior for {param_pretty}."
# #         _add_paragraph(doc, intro, italic=True)

# #         if mode == "mesh_compare":
# #             dimensions = (
# #                 curve_table[["router_key", "floor_name", "band"]]
# #                 .drop_duplicates()
# #                 .sort_values(["router_key", "floor_name", "band"])
# #                 .to_dict("records")
# #             )
# #             for row in dimensions:
# #                 router_key = clean_router_name(row["router_key"])
# #                 floor_name = normalize_floor_name(str(row["floor_name"]))
# #                 band = normalize_band_value(str(row["band"]))
# #                 subset = curve_table[
# #                     (curve_table["router_key"].astype(str) == router_key)
# #                     & (curve_table["floor_name"].astype(str) == floor_name)
# #                     & (curve_table["band"].astype(str) == band)
# #                 ]
# #                 _add_heading(doc, f"{router_key} — {floor_name} — {band}", level=2)
# #                 plot_path = _find_mesh_plot(assets["plots_actual"], router_key, floor_name, band)
# #                 if plot_path:
# #                     _add_plot(doc, plot_path)
# #                 cards = _collect_mesh_cards(with_registry, without_registry, router_key, metric_folder, floor_name, band)
# #                 _append_audit_rows(audit_rows, cards, metric_folder, floor_name, band, section_type="mesh_compare")
# #                 _add_router_cards(doc, cards)
# #                 stats_rows = []
# #                 for scenario in ["with_mesh", "without_mesh"]:
# #                     part = subset[subset["scenario"].astype(str) == scenario]
# #                     if part.empty:
# #                         continue
# #                     y_col = "p50" if "p50" in part.columns else ("mean" if "mean" in part.columns else None)
# #                     if not y_col:
# #                         continue
# #                     vals = pd.to_numeric(part[y_col], errors="coerce").dropna()
# #                     if vals.empty:
# #                         continue
# #                     stats_rows.append(
# #                         {
# #                             "router": part["scenario_label"].iloc[0] if "scenario_label" in part.columns else scenario,
# #                             "avg": float(vals.mean()),
# #                             "min": float(vals.min()),
# #                             "max": float(vals.max()),
# #                             "range": float(vals.max() - vals.min()),
# #                         }
# #                     )
# #                 stats_rows.sort(key=lambda r: r["avg"], reverse=higher_is_better)
# #                 bullets = _fallback_bullets(stats_rows, param_pretty, param_unit, floor_name, band, higher_is_better)
# #                 _add_bullets(doc, bullets)
# #                 doc.add_paragraph()
# #         else:
# #             bands = sorted(curve_table["band"].astype(str).unique())
# #             floors = sorted(curve_table["floor_name"].astype(str).unique())
# #             for band in bands:
# #                 for floor_name in floors:
# #                     subset = curve_table[
# #                         (curve_table["band"].astype(str) == band) & (curve_table["floor_name"].astype(str) == floor_name)
# #                     ]
# #                     if subset.empty:
# #                         continue
# #                     _add_heading(doc, f"{floor_name} — {band}", level=2)
# #                     plot_path = _find_rvr_plot(assets["plots_actual"], band, floor_name)
# #                     if plot_path:
# #                         _add_plot(doc, plot_path)
# #                     cards = _collect_rvr_router_cards(registry, subset, metric_folder, floor_name, band)
# #                     _append_audit_rows(audit_rows, cards, metric_folder, floor_name, band, section_type="rvr")
# #                     _add_router_cards(doc, cards)
# #                     stats_rows = _compute_stats(subset, higher_is_better)
# #                     bullets = []
# #                     if use_ai and stats_rows:
# #                         ai_text = call_ollama(_analysis_prompt(param_pretty, param_unit, floor_name, band, config_label, stats_rows, higher_is_better), model=ai_model, base_url=ai_base_url)
# #                         bullets = _parse_bullets(ai_text)
# #                     if not bullets:
# #                         bullets = _fallback_bullets(stats_rows, param_pretty, param_unit, floor_name, band, higher_is_better)
# #                     _add_bullets(doc, bullets)
# #                     doc.add_paragraph()
# #         if idx < total:
# #             doc.add_page_break()

# #     doc.save(str(output_path))
# #     _write_asset_audit(output_path, audit_rows)
# #     return output_path


# # def streamlit_report_card(current_router_dir: Optional[Path], rvr_outputs_root: Optional[Path], extracted_root: Optional[Path], step_label: str = "Step 6") -> None:
# #     import streamlit as st

# #     st.markdown('<div class="card">', unsafe_allow_html=True)
# #     st.markdown(
# #         f"""
# #         <div class="card-title">
# #           <h2><span class="step">{step_label}</span> Professional DOCX Report</h2>
# #         </div>
# #         <div class="subtle">Graph → exact matched heatmaps + exact color scales → analysis bullets. A companion asset audit CSV is written with every report.</div>
# #         """,
# #         unsafe_allow_html=True,
# #     )

# #     if current_router_dir is None:
# #         st.info("Choose or load a router folder first.")
# #         st.markdown("</div>", unsafe_allow_html=True)
# #         return

# #     mode_label = st.radio("Report mode", ["Parameter vs Range", "Mesh vs No Mesh"], horizontal=True, key="report_mode")
# #     mode = "mesh_compare" if mode_label == "Mesh vs No Mesh" else "rvr"

# #     metric_source_root = current_router_dir / ("compare_outputs" if mode == "mesh_compare" else "rvr_outputs")
# #     metric_dirs = sorted([p.name for p in metric_source_root.iterdir() if p.is_dir()]) if metric_source_root.exists() else []
# #     selected_metrics = st.multiselect("Metrics", metric_dirs, default=metric_dirs[:1], key=f"report_metrics_{mode}")
# #     config_label = st.text_input("Configuration label", value=("Mesh comparison" if mode == "mesh_compare" else "With Mesh"), key=f"report_cfg_{mode}")
# #     use_ai = st.checkbox("Use local AI analysis (Ollama)", value=False, key=f"report_use_ai_{mode}")
# #     ai_model = st.text_input("AI model", value="gemma3:4b", key=f"report_ai_model_{mode}")

# #     run_button = st.button("Generate DOCX report", key=f"report_generate_{mode}", width="stretch")
# #     if run_button:
# #         if not selected_metrics:
# #             st.error("Select at least one metric.")
# #         else:
# #             out_path = current_router_dir / ("mesh_compare_report.docx" if mode == "mesh_compare" else "comparative_report.docx")
# #             scenario_roots = None
# #             if mode == "mesh_compare":
# #                 with_root = current_router_dir / "compare_inputs" / "with_mesh_extracted"
# #                 without_root = current_router_dir / "compare_inputs" / "without_mesh_extracted"
# #                 scenario_roots = {"with_mesh": with_root, "without_mesh": without_root}
# #             try:
# #                 generate_report(
# #                     rvr_outputs_root=rvr_outputs_root or (current_router_dir / "rvr_outputs"),
# #                     extracted_root=extracted_root or (current_router_dir / "extracted"),
# #                     output_path=out_path,
# #                     metric_folders=selected_metrics,
# #                     config_label=config_label,
# #                     ai_model=ai_model,
# #                     use_ai=use_ai,
# #                     mode=mode,
# #                     compare_outputs_root=(current_router_dir / "compare_outputs"),
# #                     extracted_roots_by_scenario=scenario_roots,
# #                     csv_outputs_root=(current_router_dir / "csv_outputs"),
# #                 )
# #                 audit_path = out_path.with_name(f"{out_path.stem}_asset_audit.csv")
# #                 st.success(f"Report created: {out_path.name}")
# #                 st.download_button("Download report", data=out_path.read_bytes(), file_name=out_path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", width="stretch")
# #                 if audit_path.exists():
# #                     st.download_button("Download asset audit CSV", data=audit_path.read_bytes(), file_name=audit_path.name, mime="text/csv", width="stretch")
# #             except Exception as exc:
# #                 st.error(f"Report generation failed: {exc}")

# #     st.markdown("</div>", unsafe_allow_html=True)


# from __future__ import annotations

# import csv
# import math
# import re
# import textwrap
# from pathlib import Path
# from typing import Any, Callable, Dict, List, Optional, Tuple

# import pandas as pd
# import requests
# from docx import Document
# from docx.enum.section import WD_ORIENT
# from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.shared import Inches, Pt

# from asset_registry import AssetRegistry
# from metadata_utils import PARAM_PRETTY, canonical_metric_key, clean_router_name, normalize_band_value, normalize_floor_name

# _PARAM_UNITS = {
#     "signal_strength": "dBm",
#     "secondary_signal_strength": "dBm",
#     "tertiary_signal_strength": "dBm",
#     "snr": "dB",
#     "noise": "dBm",
#     "data_rate": "Mbps",
#     "throughput": "Mbps",
#     "channel_utilization": "%",
#     "channel_interference": "dB",
#     "channel_width": "MHz",
#     "spectrum_channel_power": "dBm",
#     "network_health": "score",
#     "network_issues": "count",
#     "number_of_access_points": "count",
# }

# _HIGHER_IS_BETTER = {
#     "signal_strength": True,
#     "secondary_signal_strength": True,
#     "tertiary_signal_strength": True,
#     "snr": True,
#     "noise": False,
#     "data_rate": True,
#     "throughput": True,
#     "channel_utilization": False,
#     "channel_interference": False,
#     "channel_width": True,
#     "spectrum_channel_power": True,
#     "network_health": True,
#     "network_issues": False,
#     "number_of_access_points": False,
# }


# def _safe_slug(text: str) -> str:
#     return re.sub(r"[^A-Za-z0-9_-]+", "_", str(text)).strip("_")


# def _fmt_value(value: float, unit: str) -> str:
#     return f"{float(value):.1f} {unit}" if unit else f"{float(value):.1f}"


# def _positive_gap(best_avg: float, worst_avg: float, higher_is_better: bool) -> float:
#     gap = float(best_avg) - float(worst_avg) if higher_is_better else float(worst_avg) - float(best_avg)
#     return abs(gap)


# def _router_list_text(rows: List[Dict[str, Any]], unit: str, *, include_max: bool = False) -> str:
#     parts: List[str] = []
#     for row in rows:
#         text = f"{row['router']} ({_fmt_value(row['avg'], unit)} avg"
#         if include_max:
#             text += f", {_fmt_value(row['max'], unit)} max"
#         text += ")"
#         parts.append(text)
#     if not parts:
#         return ""
#     if len(parts) == 1:
#         return parts[0]
#     if len(parts) == 2:
#         return f"{parts[0]} and {parts[1]}"
#     return ", ".join(parts[:-1]) + f", and {parts[-1]}"


# def _metric_specific_reason(param_key: str, band: str, higher_is_better: bool) -> str:
#     pkey = canonical_metric_key(param_key) or param_key
#     band_text = normalize_band_value(band)

#     if pkey in {"signal_strength", "secondary_signal_strength", "tertiary_signal_strength", "spectrum_channel_power"}:
#         base = "antenna pattern, receiver sensitivity, spatial-stream behavior, and attenuation handling along the surveyed path"
#     elif pkey in {"snr", "noise", "channel_interference", "channel_utilization"}:
#         base = "co-channel activity, noise floor control, interference rejection, and airtime management"
#     elif pkey in {"throughput", "data_rate"}:
#         base = "modulation-and-coding stability, retry behavior, airtime efficiency, and backhaul effectiveness under load"
#     elif pkey in {"network_health", "network_issues"}:
#         base = "link stability, retransmission behavior, and the way the platform manages contention and roaming events"
#     elif pkey in {"number_of_access_points"}:
#         base = "cell planning, roaming decisions, and how aggressively the client transitions between available access points"
#     else:
#         base = "radio-chain performance, propagation loss, and how efficiently the platform handles interference and airtime"

#     if band_text.startswith("6"):
#         band_note = " At 6 GHz, the separation is usually amplified by the higher free-space loss and weaker wall penetration."
#     elif band_text.startswith("5"):
#         band_note = " On 5 GHz, the separation often reflects how well each router preserves link quality as attenuation rises with distance."
#     elif band_text.startswith("2.4"):
#         band_note = " On 2.4 GHz, congestion tolerance and interference handling usually play a larger role than pure path loss alone."
#     else:
#         band_note = ""

#     direction = "higher" if higher_is_better else "lower"
#     return f"The separation is most plausibly driven by differences in {base}; for this metric, {direction} values indicate the stronger result.{band_note}"


# def _get_param_info(param_key: str) -> Tuple[str, str, bool]:
#     pkey = canonical_metric_key(param_key) or param_key
#     return (
#         PARAM_PRETTY.get(pkey, pkey.replace("_", " ").title()),
#         _PARAM_UNITS.get(pkey, ""),
#         _HIGHER_IS_BETTER.get(pkey, True),
#     )


# def call_ollama(prompt: str, model: str = "gemma3:4b", base_url: str = "http://localhost:11434", timeout: int = 90) -> str:
#     errors = []

#     try:
#         resp = requests.post(
#             f"{base_url.rstrip('/')}/api/generate",
#             json={
#                 "model": model,
#                 "prompt": prompt,
#                 "stream": False,
#                 "options": {"temperature": 0.2, "top_p": 0.9, "num_predict": 220},
#             },
#             timeout=timeout,
#         )
#         if resp.status_code == 404:
#             raise requests.exceptions.HTTPError(f"404 at {base_url.rstrip('/')}/api/generate")
#         resp.raise_for_status()
#         body = resp.json()
#         if isinstance(body, dict) and "response" in body and body.get("response"):
#             return _strip_preamble(str(body.get("response", "")).strip())
#     except Exception as exc:
#         errors.append(f"legacy /api/generate: {exc}")

#     try:
#         resp = requests.post(
#             f"{base_url.rstrip('/')}/v1/chat/completions",
#             json={
#                 "model": model,
#                 "messages": [{"role": "user", "content": prompt}],
#                 "temperature": 0.2,
#                 "max_tokens": 220,
#             },
#             timeout=timeout,
#         )
#         resp.raise_for_status()
#         body = resp.json()
#         choices = body.get("choices", [])
#         if choices:
#             content = choices[0].get("message", {}).get("content")
#             if content:
#                 return _strip_preamble(str(content).strip())
#     except Exception as exc:
#         errors.append(f"v1/chat/completions: {exc}")

#     try:
#         resp = requests.post(
#             f"{base_url.rstrip('/')}/v1/completions",
#             json={
#                 "model": model,
#                 "prompt": prompt,
#                 "temperature": 0.2,
#                 "max_tokens": 220,
#             },
#             timeout=timeout,
#         )
#         resp.raise_for_status()
#         body = resp.json()
#         choices = body.get("choices", [])
#         if choices:
#             text = choices[0].get("text") or choices[0].get("message", {}).get("content")
#             if text:
#                 return _strip_preamble(str(text).strip())
#     except Exception as exc:
#         errors.append(f"v1/completions: {exc}")

#     return f"[AI: {' | '.join(errors)}]"


# def _strip_preamble(text: str) -> str:
#     if not text:
#         return text
#     text = text.strip()
#     preamble_patterns = [
#         r"^Here'?s?\s+(?:a\s+)?(?:comparative\s+)?(?:report|summary|analysis)[^:]*:\s*",
#         r"^Based on\s+the\s+provided\s+data[^:]*:\s*",
#         r"^Here are\s+",
#         r"^Here'?s?\s+",
#     ]
#     for pattern in preamble_patterns:
#         text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
#     conclusion_patterns = [
#         r"\n*(?:This|These|The above)\s+(?:analysis|summary|points?|bullets?).*$",
#         r"\n*(?:Overall|In summary|In conclusion)[^.]*$",
#         r"\n*Feel free to.*$",
#     ]
#     for pattern in conclusion_patterns:
#         text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
#     return text.strip()


# def _parse_bullets(text: str) -> List[str]:
#     text = _strip_preamble(text)
#     lines: List[str] = []
#     for line in (text or "").splitlines():
#         line = line.strip()
#         if not line:
#             continue
#         line = re.sub(r"^[\-•*\d\.\)\s]+", "", line)
#         line = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", line)
#         line = line.strip()
#         if len(line) >= 10:
#             lines.append(line)
#     return lines


# def _technical_reason_prompt(param_pretty: str, param_unit: str, floor_name: str, band: str, stats_rows: List[Dict[str, Any]], higher_is_better: bool) -> str:
#     direction_word = "higher values are stronger" if higher_is_better else "lower values are stronger"
#     rows_text = "\n".join(
#         f"- {row['router']}: avg={row['avg']:.1f} {param_unit}, max={row['max']:.1f} {param_unit}, min={row['min']:.1f} {param_unit}"
#         for row in stats_rows
#     )
#     return textwrap.dedent(
#         f"""
#         Write exactly one bullet line that starts with •.
#         Explain the likely technical reason for the separation in a Wi-Fi comparative graph.
#         Stay grounded in radio/network science. Do not change the ranking or mention any unsupported facts.

#         Metric: {param_pretty} ({param_unit})
#         Floor: {floor_name}
#         Band: {band}
#         Interpretation: {direction_word}

#         Router statistics:
#         {rows_text}
#         """
#     ).strip()


# def _build_summary_bullets(
#     stats_rows: List[Dict[str, Any]],
#     param_key: str,
#     param_pretty: str,
#     param_unit: str,
#     floor_name: str,
#     band: str,
#     higher_is_better: bool,
#     *,
#     use_ai: bool = False,
#     ai_model: str = "gemma3:4b",
#     ai_base_url: str = "http://localhost:11434",
# ) -> List[str]:
#     if not stats_rows:
#         return [f"No usable data was available for {floor_name} on {band}."]

#     best = stats_rows[0]
#     worst = stats_rows[-1]
#     mid_rows = stats_rows[1:-1] if len(stats_rows) >= 3 else []
#     best_to_worst_gap = _positive_gap(best["avg"], worst["avg"], higher_is_better)
#     within_router_ranges = [float(row["range"]) for row in stats_rows]
#     min_range = min(within_router_ranges)
#     max_range = max(within_router_ranges)

#     if len(stats_rows) == 1:
#         pattern_text = (
#             f"The comparative study graph for {floor_name} on {band} contains only one router trace, so it shows a single baseline with "
#             f"{_fmt_value(best['range'], param_unit)} of within-curve spread across distance bins."
#         )
#     else:
#         pattern_text = (
#             f"The comparative study graph for {floor_name} on {band} shows a clear best-to-worst average separation of "
#             f"{_fmt_value(best_to_worst_gap, param_unit)}, while within-router variability spans "
#             f"{_fmt_value(min_range, param_unit)} to {_fmt_value(max_range, param_unit)} across the plotted distance bins."
#         )

#     bullet_1 = pattern_text
#     bullet_2 = (
#         f"{best['router']} is the best performer with an average of {_fmt_value(best['avg'], param_unit)} "
#         f"and a maximum of {_fmt_value(best['max'], param_unit)}."
#     )

#     if mid_rows:
#         bullet_3 = f"The mid-tier router set is {_router_list_text(mid_rows, param_unit, include_max=True)}."
#     else:
#         bullet_3 = "No mid-tier router classification applies here because fewer than three routers are present in this floor-and-band comparison."

#     bullet_4 = (
#         f"{worst['router']} is the weakest performer with an average of {_fmt_value(worst['avg'], param_unit)} "
#         f"and a maximum of {_fmt_value(worst['max'], param_unit)}; the best-to-worst average gap is {_fmt_value(best_to_worst_gap, param_unit)}."
#     )

#     bullet_5 = _metric_specific_reason(param_key, band, higher_is_better)
#     if use_ai and stats_rows:
#         try:
#             ai_text = call_ollama(
#                 _technical_reason_prompt(param_pretty, param_unit, floor_name, band, stats_rows, higher_is_better),
#                 model=ai_model,
#                 base_url=ai_base_url,
#             )
#             ai_lines = _parse_bullets(ai_text)
#             if ai_lines:
#                 bullet_5 = ai_lines[0]
#         except Exception:
#             pass

#     return [bullet_1, bullet_2, bullet_3, bullet_4, bullet_5]


# def discover_rvr_assets(root: Path, metric_folder: str) -> Dict[str, Any]:
#     metric_dir = root / metric_folder
#     plots_actual: Dict[Tuple[str, str], Path] = {}
#     table_path = None
#     tables_dir = metric_dir / "tables"
#     if tables_dir.exists():
#         candidates = sorted(tables_dir.glob("*_curve_tables.csv"))
#         if candidates:
#             table_path = candidates[0]
#     plot_dir = metric_dir / "plots_actual" / "COMPARE_ROUTERS_BY_FLOOR"
#     if plot_dir.exists():
#         for png in sorted(plot_dir.glob("*.png")):
#             m = re.match(r"^.+?__act__(?P<band>.+?)__floor_(?P<floor>.+?)\.png$", png.name)
#             if m:
#                 plots_actual[(m.group("band"), m.group("floor"))] = png
#     return {"mode": "rvr", "metric_dir": metric_dir, "curve_table_path": table_path, "plots_actual": plots_actual}


# def discover_mesh_assets(root: Path, metric_folder: str) -> Dict[str, Any]:
#     metric_dir = root / metric_folder
#     plot_map: Dict[Tuple[str, str, str], Path] = {}
#     table_path = None
#     tables_dir = metric_dir / "tables"
#     if tables_dir.exists():
#         candidates = sorted(tables_dir.glob("*_mesh_curve_tables.csv"))
#         if candidates:
#             table_path = candidates[0]
#     for png in sorted(metric_dir.rglob("*.png")):
#         if png.parent == metric_dir:
#             continue
#         router = clean_router_name(png.parent.parent.name) if png.parent.parent != metric_dir else ""
#         floor_name = normalize_floor_name(png.parent.name)
#         stem = png.stem
#         if "_" not in stem:
#             continue
#         band, _ = stem.split("_", 1)
#         plot_map[(router, floor_name, normalize_band_value(band))] = png
#     return {"mode": "mesh_compare", "metric_dir": metric_dir, "curve_table_path": table_path, "plots_actual": plot_map}


# def _find_rvr_plot(plot_map: Dict[Tuple[str, str], Path], band: str, floor_name: str) -> Optional[Path]:
#     band = normalize_band_value(band)
#     floor_norm = _safe_slug(normalize_floor_name(floor_name))
#     key = (band, floor_norm)
#     if key in plot_map:
#         return plot_map[key]
#     for (bs, fs), path in plot_map.items():
#         if band in bs and floor_norm.lower() in fs.lower():
#             return path
#     return None


# def _find_mesh_plot(plot_map: Dict[Tuple[str, str, str], Path], router_key: str, floor_name: str, band: str) -> Optional[Path]:
#     key = (clean_router_name(router_key), normalize_floor_name(floor_name), normalize_band_value(band))
#     if key in plot_map:
#         return plot_map[key]
#     return None


# def _pick_y_col(subset: pd.DataFrame) -> Optional[str]:
#     for col in ["p50", "mean", "p90", "p10", "max", "min"]:
#         if col in subset.columns:
#             return col
#     return None


# def _compute_stats(subset: pd.DataFrame, higher_is_better: bool) -> List[Dict[str, Any]]:
#     y_col = _pick_y_col(subset)
#     if not y_col:
#         return []

#     rows: List[Dict[str, Any]] = []
#     for router_key in subset["router_key"].astype(str).unique():
#         part = subset[subset["router_key"].astype(str) == router_key].copy()
#         s = pd.to_numeric(part[y_col], errors="coerce").dropna()
#         if s.empty:
#             continue

#         display_series = part["router_display"] if "router_display" in part.columns else pd.Series([router_key])
#         display = clean_router_name(str(display_series.iloc[0]))
#         rows.append(
#             {
#                 "router": display,
#                 "router_key": clean_router_name(router_key),
#                 "y_col": y_col,
#                 "avg": float(s.mean()),
#                 "min": float(s.min()),
#                 "max": float(s.max()),
#                 "range": float(s.max() - s.min()),
#                 "points": int(s.shape[0]),
#             }
#         )

#     rows.sort(
#         key=lambda row: (
#             row["avg"] if higher_is_better else -row["avg"],
#             row["max"] if higher_is_better else -row["max"],
#             -row["range"],
#             row["router"].lower(),
#         ),
#         reverse=True,
#     )
#     return rows


# def _set_landscape(section) -> None:
#     section.orientation = WD_ORIENT.LANDSCAPE
#     section.page_width, section.page_height = section.page_height, section.page_width
#     section.left_margin = Inches(0.6)
#     section.right_margin = Inches(0.6)
#     section.top_margin = Inches(0.7)
#     section.bottom_margin = Inches(0.7)


# def _set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
#     tc = cell._tc
#     tcPr = tc.get_or_add_tcPr()
#     tcMar = tcPr.first_child_found_in("w:tcMar")
#     if tcMar is None:
#         tcMar = OxmlElement("w:tcMar")
#         tcPr.append(tcMar)
#     for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
#         element = tcMar.find(qn(f"w:{tag}"))
#         if element is None:
#             element = OxmlElement(f"w:{tag}")
#             tcMar.append(element)
#         element.set(qn("w:w"), str(value))
#         element.set(qn("w:type"), "dxa")


# def _add_heading(doc: Document, text: str, level: int = 1) -> None:
#     p = doc.add_heading(level=level)
#     run = p.add_run(text)
#     run.bold = True


# def _add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
#     p = doc.add_paragraph()
#     run = p.add_run(text)
#     run.italic = italic
#     run.font.size = Pt(10.5)


# def _add_plot(doc: Document, plot_path: Path) -> None:
#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     p.add_run().add_picture(str(plot_path), width=Inches(9.2))


# def _add_router_cards(doc: Document, router_cards: List[Dict[str, Any]]) -> None:
#     if not router_cards:
#         return
#     cols = 2
#     rows = math.ceil(len(router_cards) / cols)
#     table = doc.add_table(rows=rows, cols=cols)
#     table.alignment = WD_TABLE_ALIGNMENT.CENTER
#     table.style = "Table Grid"
#     table.autofit = False
#     card_width = 4.85
#     heatmap_width = 4.35
#     scale_width = 4.0

#     for idx, card in enumerate(router_cards):
#         cell = table.cell(idx // cols, idx % cols)
#         cell.width = Inches(card_width)
#         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#         _set_cell_margins(cell, top=90, start=90, bottom=90, end=90)

#         p = cell.paragraphs[0]
#         p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         name_run = p.add_run(card["display"])
#         name_run.bold = True
#         name_run.font.size = Pt(10)

#         if card.get("scenario_label"):
#             label_p = cell.add_paragraph()
#             label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             label_run = label_p.add_run(card["scenario_label"])
#             label_run.italic = True
#             label_run.font.size = Pt(9)

#         if card.get("heatmap") and Path(card["heatmap"]).exists():
#             hp = cell.add_paragraph()
#             hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             hp.add_run().add_picture(str(card["heatmap"]), width=Inches(heatmap_width))
#         else:
#             miss = cell.add_paragraph()
#             miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             miss.add_run("Exact heatmap match not found").italic = True

#         if card.get("scale") and Path(card["scale"]).exists():
#             sp = cell.add_paragraph()
#             sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             sp.add_run().add_picture(str(card["scale"]), width=Inches(scale_width))
#         else:
#             miss = cell.add_paragraph()
#             miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             miss.add_run("Matching color scale not found").italic = True

#     total_cells = rows * cols
#     for idx in range(len(router_cards), total_cells):
#         table.cell(idx // cols, idx % cols).text = ""


# def _add_bullets(doc: Document, bullets: List[str]) -> None:
#     for bullet in bullets:
#         p = doc.add_paragraph(style="List Bullet")
#         run = p.add_run(bullet)
#         run.font.size = Pt(10.5)


# def _collect_rvr_router_cards(registry: AssetRegistry, subset: pd.DataFrame, parameter_key: str, floor_name: str, band: str) -> List[Dict[str, Any]]:
#     cards: List[Dict[str, Any]] = []
#     for router_key in subset["router_key"].astype(str).unique():
#         part = subset[subset["router_key"].astype(str) == router_key]
#         display = clean_router_name(str(part["router_display"].iloc[0] if "router_display" in part.columns else router_key))
#         heatmap, scale = registry.get_pair(router_key, parameter_key, floor_name, band)
#         cards.append({"router_key": clean_router_name(router_key), "display": display, "heatmap": heatmap, "scale": scale})
#     return cards


# def _collect_mesh_cards(with_registry: AssetRegistry, without_registry: AssetRegistry, router_key: str, parameter_key: str, floor_name: str, band: str) -> List[Dict[str, Any]]:
#     display = clean_router_name(router_key)
#     with_heatmap, with_scale = with_registry.get_pair(router_key, parameter_key, floor_name, band)
#     wo_heatmap, wo_scale = without_registry.get_pair(router_key, parameter_key, floor_name, band)
#     return [
#         {"router_key": display, "display": display, "scenario_label": "With mesh", "heatmap": with_heatmap, "scale": with_scale},
#         {"router_key": display, "display": display, "scenario_label": "Without mesh", "heatmap": wo_heatmap, "scale": wo_scale},
#     ]


# def _append_audit_rows(audit_rows: List[Dict[str, Any]], cards: List[Dict[str, Any]], parameter_key: str, floor_name: str, band: str, section_type: str):
#     for card in cards:
#         heatmap = card.get("heatmap")
#         scale = card.get("scale")
#         status = "exact_pair"
#         if not heatmap and not scale:
#             status = "missing_heatmap_and_scale"
#         elif not heatmap:
#             status = "missing_heatmap"
#         elif not scale:
#             status = "missing_scale"
#         audit_rows.append(
#             {
#                 "section_type": section_type,
#                 "router_key": card.get("router_key", ""),
#                 "display": card.get("display", ""),
#                 "scenario_label": card.get("scenario_label", ""),
#                 "parameter_key": canonical_metric_key(parameter_key) or parameter_key,
#                 "floor_name": normalize_floor_name(floor_name),
#                 "band": normalize_band_value(band),
#                 "heatmap_path": str(heatmap) if heatmap else "",
#                 "scale_path": str(scale) if scale else "",
#                 "status": status,
#             }
#         )


# def _write_asset_audit(output_path: Path, audit_rows: List[Dict[str, Any]]) -> Optional[Path]:
#     if not audit_rows:
#         return None
#     audit_path = output_path.with_name(f"{output_path.stem}_asset_audit.csv")
#     with audit_path.open("w", encoding="utf-8", newline="") as handle:
#         writer = csv.DictWriter(handle, fieldnames=list(audit_rows[0].keys()))
#         writer.writeheader()
#         writer.writerows(audit_rows)
#     return audit_path


# def generate_report(
#     rvr_outputs_root: Path,
#     extracted_root: Path,
#     output_path: Path,
#     metric_folders: List[str],
#     config_label: str = "Standard",
#     ai_model: str = "gemma3:4b",
#     ai_base_url: str = "http://localhost:11434",
#     use_ai: bool = True,
#     progress_cb: Optional[Callable[[int, int, str], None]] = None,
#     mode: str = "rvr",
#     compare_outputs_root: Optional[Path] = None,
#     extracted_roots_by_scenario: Optional[Dict[str, Path]] = None,
#     csv_outputs_root: Optional[Path] = None,
# ) -> Path:
#     doc = Document()
#     _set_landscape(doc.sections[0])
#     title = doc.add_paragraph()
#     title.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     run = title.add_run("Wi-Fi Comparative Analysis Report")
#     run.bold = True
#     run.font.size = Pt(18)
#     subtitle = "Mesh vs No Mesh" if mode == "mesh_compare" else "Router-to-router comparison"
#     _add_paragraph(doc, subtitle, italic=True)

#     total = max(1, len(metric_folders))
#     registry = AssetRegistry.from_roots(extracted_root=extracted_root, csv_outputs_root=csv_outputs_root)
#     with_registry = without_registry = AssetRegistry([])
#     if mode == "mesh_compare" and extracted_roots_by_scenario:
#         with_registry = AssetRegistry.from_roots(extracted_root=extracted_roots_by_scenario.get("with_mesh", Path(".")))
#         without_registry = AssetRegistry.from_roots(extracted_root=extracted_roots_by_scenario.get("without_mesh", Path(".")))

#     audit_rows: List[Dict[str, Any]] = []

#     for idx, metric_folder in enumerate(metric_folders, start=1):
#         metric_folder = canonical_metric_key(metric_folder) or metric_folder
#         if progress_cb:
#             progress_cb(idx, total, f"Preparing {metric_folder}")
#         param_pretty, param_unit, higher_is_better = _get_param_info(metric_folder)
#         heading_suffix = f" — {config_label}" if config_label and config_label.strip().lower() not in {"", "standard"} else ""
#         _add_heading(doc, f"Range vs {param_pretty}{heading_suffix}", level=1)

#         if mode == "mesh_compare":
#             if compare_outputs_root is None:
#                 raise ValueError("compare_outputs_root is required for mesh_compare mode")
#             assets = discover_mesh_assets(compare_outputs_root, metric_folder)
#         else:
#             assets = discover_rvr_assets(rvr_outputs_root, metric_folder)

#         curve_table_path = assets.get("curve_table_path")
#         if not curve_table_path or not Path(curve_table_path).exists():
#             _add_paragraph(doc, f"No curve table was found for {param_pretty}.")
#             continue

#         curve_table = pd.read_csv(curve_table_path)
#         if curve_table.empty:
#             _add_paragraph(doc, f"The curve table for {param_pretty} is empty.")
#             continue
#         if "band" in curve_table.columns:
#             curve_table["band"] = curve_table["band"].astype(str).map(normalize_band_value)
#         if "floor_name" in curve_table.columns:
#             curve_table["floor_name"] = curve_table["floor_name"].astype(str).map(normalize_floor_name)
#         if "router_key" in curve_table.columns:
#             curve_table["router_key"] = curve_table["router_key"].astype(str).map(clean_router_name)
#         if "router_display" in curve_table.columns:
#             curve_table["router_display"] = curve_table["router_display"].astype(str).map(clean_router_name)

#         intro = f"This section summarizes the comparative-study graph for {param_pretty} using the curve-table statistics." if mode == "rvr" else f"This section compares with-mesh and without-mesh behavior for {param_pretty}."
#         _add_paragraph(doc, intro, italic=True)

#         if mode == "mesh_compare":
#             dimensions = (
#                 curve_table[["router_key", "floor_name", "band"]]
#                 .drop_duplicates()
#                 .sort_values(["router_key", "floor_name", "band"])
#                 .to_dict("records")
#             )
#             for row in dimensions:
#                 router_key = clean_router_name(row["router_key"])
#                 floor_name = normalize_floor_name(str(row["floor_name"]))
#                 band = normalize_band_value(str(row["band"]))
#                 subset = curve_table[
#                     (curve_table["router_key"].astype(str) == router_key)
#                     & (curve_table["floor_name"].astype(str) == floor_name)
#                     & (curve_table["band"].astype(str) == band)
#                 ]
#                 _add_heading(doc, f"{router_key} — {floor_name} — {band}", level=2)
#                 plot_path = _find_mesh_plot(assets["plots_actual"], router_key, floor_name, band)
#                 if plot_path:
#                     _add_plot(doc, plot_path)
#                 cards = _collect_mesh_cards(with_registry, without_registry, router_key, metric_folder, floor_name, band)
#                 _append_audit_rows(audit_rows, cards, metric_folder, floor_name, band, section_type="mesh_compare")
#                 _add_router_cards(doc, cards)
#                 stats_rows = []
#                 for scenario in ["with_mesh", "without_mesh"]:
#                     part = subset[subset["scenario"].astype(str) == scenario]
#                     if part.empty:
#                         continue
#                     y_col = _pick_y_col(part)
#                     if not y_col:
#                         continue
#                     vals = pd.to_numeric(part[y_col], errors="coerce").dropna()
#                     if vals.empty:
#                         continue
#                     stats_rows.append(
#                         {
#                             "router": part["scenario_label"].iloc[0] if "scenario_label" in part.columns else scenario,
#                             "avg": float(vals.mean()),
#                             "min": float(vals.min()),
#                             "max": float(vals.max()),
#                             "range": float(vals.max() - vals.min()),
#                         }
#                     )
#                 stats_rows.sort(key=lambda row: row["avg"], reverse=higher_is_better)
#                 bullets = _build_summary_bullets(
#                     stats_rows,
#                     metric_folder,
#                     param_pretty,
#                     param_unit,
#                     floor_name,
#                     band,
#                     higher_is_better,
#                     use_ai=use_ai,
#                     ai_model=ai_model,
#                     ai_base_url=ai_base_url,
#                 )
#                 _add_bullets(doc, bullets)
#                 doc.add_paragraph()
#         else:
#             bands = sorted(curve_table["band"].astype(str).unique())
#             floors = sorted(curve_table["floor_name"].astype(str).unique())
#             for band in bands:
#                 for floor_name in floors:
#                     subset = curve_table[
#                         (curve_table["band"].astype(str) == band) & (curve_table["floor_name"].astype(str) == floor_name)
#                     ]
#                     if subset.empty:
#                         continue
#                     _add_heading(doc, f"{floor_name} — {band}", level=2)
#                     plot_path = _find_rvr_plot(assets["plots_actual"], band, floor_name)
#                     if plot_path:
#                         _add_plot(doc, plot_path)
#                     cards = _collect_rvr_router_cards(registry, subset, metric_folder, floor_name, band)
#                     _append_audit_rows(audit_rows, cards, metric_folder, floor_name, band, section_type="rvr")
#                     _add_router_cards(doc, cards)
#                     stats_rows = _compute_stats(subset, higher_is_better)
#                     bullets = _build_summary_bullets(
#                         stats_rows,
#                         metric_folder,
#                         param_pretty,
#                         param_unit,
#                         floor_name,
#                         band,
#                         higher_is_better,
#                         use_ai=use_ai,
#                         ai_model=ai_model,
#                         ai_base_url=ai_base_url,
#                     )
#                     _add_bullets(doc, bullets)
#                     doc.add_paragraph()
#         if idx < total:
#             doc.add_page_break()

#     doc.save(str(output_path))
#     _write_asset_audit(output_path, audit_rows)
#     return output_path


# def streamlit_report_card(current_router_dir: Optional[Path], rvr_outputs_root: Optional[Path], extracted_root: Optional[Path], step_label: str = "Step 6") -> None:
#     import streamlit as st

#     st.markdown('<div class="card">', unsafe_allow_html=True)
#     st.markdown(
#         f"""
#         <div class="card-title">
#           <h2><span class="step">{step_label}</span> Professional DOCX Report</h2>
#         </div>
#         <div class="subtle">Graph → exact matched heatmaps + exact color scales → graph-aligned summary bullets. A companion asset audit CSV is written with every report.</div>
#         """,
#         unsafe_allow_html=True,
#     )

#     if current_router_dir is None:
#         st.info("Choose or load a router folder first.")
#         st.markdown("</div>", unsafe_allow_html=True)
#         return

#     mode_label = st.radio("Report mode", ["Parameter vs Range", "Mesh vs No Mesh"], horizontal=True, key="report_mode")
#     mode = "mesh_compare" if mode_label == "Mesh vs No Mesh" else "rvr"

#     metric_source_root = current_router_dir / ("compare_outputs" if mode == "mesh_compare" else "rvr_outputs")
#     metric_dirs = sorted([p.name for p in metric_source_root.iterdir() if p.is_dir()]) if metric_source_root.exists() else []
#     selected_metrics = st.multiselect("Metrics", metric_dirs, default=metric_dirs[:1], key=f"report_metrics_{mode}")
#     config_label = st.text_input("Configuration label", value=("Mesh comparison" if mode == "mesh_compare" else "With Mesh"), key=f"report_cfg_{mode}")
#     use_ai = st.checkbox("Use local AI analysis (Ollama)", value=False, key=f"report_use_ai_{mode}")
#     ai_model = st.text_input("AI model", value="gemma3:4b", key=f"report_ai_model_{mode}")

#     run_button = st.button("Generate DOCX report", key=f"report_generate_{mode}", width="stretch")
#     if run_button:
#         if not selected_metrics:
#             st.error("Select at least one metric.")
#         else:
#             out_path = current_router_dir / ("mesh_compare_report.docx" if mode == "mesh_compare" else "comparative_report.docx")
#             scenario_roots = None
#             if mode == "mesh_compare":
#                 with_root = current_router_dir / "compare_inputs" / "with_mesh_extracted"
#                 without_root = current_router_dir / "compare_inputs" / "without_mesh_extracted"
#                 scenario_roots = {"with_mesh": with_root, "without_mesh": without_root}
#             try:
#                 generate_report(
#                     rvr_outputs_root=rvr_outputs_root or (current_router_dir / "rvr_outputs"),
#                     extracted_root=extracted_root or (current_router_dir / "extracted"),
#                     output_path=out_path,
#                     metric_folders=selected_metrics,
#                     config_label=config_label,
#                     ai_model=ai_model,
#                     use_ai=use_ai,
#                     mode=mode,
#                     compare_outputs_root=(current_router_dir / "compare_outputs"),
#                     extracted_roots_by_scenario=scenario_roots,
#                     csv_outputs_root=(current_router_dir / "csv_outputs"),
#                 )
#                 audit_path = out_path.with_name(f"{out_path.stem}_asset_audit.csv")
#                 st.success(f"Report created: {out_path.name}")
#                 st.download_button("Download report", data=out_path.read_bytes(), file_name=out_path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", width="stretch")
#                 if audit_path.exists():
#                     st.download_button("Download asset audit CSV", data=audit_path.read_bytes(), file_name=audit_path.name, mime="text/csv", width="stretch")
#             except Exception as exc:
#                 st.error(f"Report generation failed: {exc}")

#     st.markdown("</div>", unsafe_allow_html=True)


# from __future__ import annotations

# import csv
# import math
# import re
# import textwrap
# from pathlib import Path
# from typing import Any, Callable, Dict, List, Optional, Tuple

# import pandas as pd
# import requests
# from docx import Document
# from docx.enum.section import WD_ORIENT
# from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.shared import Inches, Pt

# from asset_registry import AssetRegistry
# from metadata_utils import PARAM_PRETTY, canonical_metric_key, clean_router_name, normalize_band_value, normalize_floor_name

# _PARAM_UNITS = {
#     "signal_strength": "dBm",
#     "secondary_signal_strength": "dBm",
#     "tertiary_signal_strength": "dBm",
#     "snr": "dB",
#     "noise": "dBm",
#     "data_rate": "Mbps",
#     "throughput": "Mbps",
#     "channel_utilization": "%",
#     "channel_interference": "dB",
#     "channel_width": "MHz",
#     "spectrum_channel_power": "dBm",
#     "network_health": "score",
#     "network_issues": "count",
#     "number_of_access_points": "count",
# }

# _HIGHER_IS_BETTER = {
#     "signal_strength": True,
#     "secondary_signal_strength": True,
#     "tertiary_signal_strength": True,
#     "snr": True,
#     "noise": False,
#     "data_rate": True,
#     "throughput": True,
#     "channel_utilization": False,
#     "channel_interference": False,
#     "channel_width": True,
#     "spectrum_channel_power": True,
#     "network_health": True,
#     "network_issues": False,
#     "number_of_access_points": False,
# }


# def _safe_slug(text: str) -> str:
#     return re.sub(r"[^A-Za-z0-9_-]+", "_", str(text)).strip("_")


# def _get_param_info(param_key: str) -> Tuple[str, str, bool]:
#     pkey = canonical_metric_key(param_key) or param_key
#     return (
#         PARAM_PRETTY.get(pkey, pkey.replace("_", " ").title()),
#         _PARAM_UNITS.get(pkey, ""),
#         _HIGHER_IS_BETTER.get(pkey, True),
#     )


# def call_ollama(prompt: str, model: str = "gemma3:4b", base_url: str = "http://localhost:11434", timeout: int = 90) -> str:
#     errors = []

#     # 1) Ollama legacy endpoint (/api/generate)
#     try:
#         resp = requests.post(
#             f"{base_url.rstrip('/')}/api/generate",
#             json={
#                 "model": model,
#                 "prompt": prompt,
#                 "stream": False,
#                 "options": {"temperature": 0.2, "top_p": 0.9, "num_predict": 700},
#             },
#             timeout=timeout,
#         )
#         if resp.status_code == 404:
#             raise requests.exceptions.HTTPError(f"404 at {base_url.rstrip('/')}/api/generate")
#         resp.raise_for_status()
#         body = resp.json()
#         if isinstance(body, dict) and "response" in body and body.get("response"):
#             return _strip_preamble(str(body.get("response", "")).strip())
#     except Exception as exc:
#         errors.append(f"legacy /api/generate: {exc}")

#     # 2) Ollama v1 chat completion endpoint (/v1/chat/completions)
#     try:
#         resp = requests.post(
#             f"{base_url.rstrip('/')}/v1/chat/completions",
#             json={
#                 "model": model,
#                 "messages": [{"role": "user", "content": prompt}],
#                 "temperature": 0.2,
#                 "max_tokens": 700,
#             },
#             timeout=timeout,
#         )
#         resp.raise_for_status()
#         body = resp.json()
#         choices = body.get("choices", [])
#         if choices:
#             content = choices[0].get("message", {}).get("content")
#             if content:
#                 return _strip_preamble(str(content).strip())
#     except Exception as exc:
#         errors.append(f"v1/chat/completions: {exc}")

#     # 3) OpenAI-compatible completions endpoint (/v1/completions)
#     try:
#         resp = requests.post(
#             f"{base_url.rstrip('/')}/v1/completions",
#             json={
#                 "model": model,
#                 "prompt": prompt,
#                 "temperature": 0.2,
#                 "max_tokens": 700,
#             },
#             timeout=timeout,
#         )
#         resp.raise_for_status()
#         body = resp.json()
#         choices = body.get("choices", [])
#         if choices:
#             text = choices[0].get("text") or choices[0].get("message", {}).get("content")
#             if text:
#                 return _strip_preamble(str(text).strip())
#     except Exception as exc:
#         errors.append(f"v1/completions: {exc}")

#     return f"[AI: {' | '.join(errors)}]"


# def _strip_preamble(text: str) -> str:
#     """Remove common AI preambles and conclusions from response."""
#     if not text:
#         return text
#     text = text.strip()
#     # Remove leading preamble phrases
#     preamble_patterns = [
#         r"^Here'?s?\s+(?:a\s+)?(?:comparative\s+)?(?:report|summary|analysis)[^:]*:\s*",
#         r"^Based on\s+the\s+provided\s+data[^:]*:\s*",
#         r"^Here are\s+",
#         r"^Here'?s?\s+",
#     ]
#     for pattern in preamble_patterns:
#         text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
#     text = text.strip()
    
#     # Remove trailing conclusion phrases
#     conclusion_patterns = [
#         r"\n*(?:This|These|The above)\s+(?:analysis|summary|points?|bullets?).*$",
#         r"\n*(?:Overall|In summary|In conclusion)[^.]*$",
#         r"\n*Feel free to.*$",
#     ]
#     for pattern in conclusion_patterns:
#         text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
#     return text.strip()


# def _parse_bullets(text: str) -> List[str]:
#     text = _strip_preamble(text)
#     lines: List[str] = []
#     for line in (text or "").splitlines():
#         line = line.strip()
#         if not line:
#             continue
#         # Remove leading bullet markers
#         line = re.sub(r"^[\-•*\d\.\)\s]+", "", line)
#         # Remove markdown bold/italic markers
#         line = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", line)
#         line = line.strip()
#         if len(line) >= 10:
#             lines.append(line)
#     return lines


# def _fallback_bullets(stats_rows: List[Dict[str, Any]], param_pretty: str, param_unit: str, floor_name: str, band: str, higher_is_better: bool) -> List[str]:
#     if not stats_rows:
#         return [f"No usable data was available for {floor_name} on {band}."]

#     best = stats_rows[0]
#     worst = stats_rows[-1]
#     spread = best["avg"] - worst["avg"]
#     ranking = " > ".join(r["router"] for r in stats_rows)
#     trend = "more stable" if max(r["range"] for r in stats_rows) < max(abs(best["avg"]) * 0.15, 1) else "more variable"
#     cluster = None
#     if len(stats_rows) >= 3:
#         mids = stats_rows[1:-1]
#         cluster = ", ".join(f"{r['router']} ({r['avg']:.1f} {param_unit})" for r in mids)

#     direction = "higher" if higher_is_better else "lower"
#     impairment = "attenuation, interference, or backhaul efficiency" if str(band).startswith("5") else "radio design, sensitivity, or congestion handling"
#     bullets = [
#         f"Across {floor_name} on {band}, the curves are {trend} than in the strongest cases, so router design matters more than distance alone.",
#         f"{best['router']} delivers the best average {param_pretty} at {best['avg']:.1f} {param_unit}.",
#     ]
#     if cluster:
#         bullets.append(f"The middle tier is formed by {cluster}.")
#     bullets.extend(
#         [
#             f"{worst['router']} is the weakest result at {worst['avg']:.1f} {param_unit}, creating a best-to-worst gap of {spread:.1f} {param_unit}.",
#             f"The separation likely reflects differences in {impairment}, especially where {direction} {param_pretty.lower()} is preferred.",
#             f"Overall ranking: {ranking}.",
#         ]
#     )
#     return bullets


# def _analysis_prompt(param_pretty: str, param_unit: str, floor_name: str, band: str, config_label: str, stats_rows: List[Dict[str, Any]], higher_is_better: bool) -> str:
#     direction_word = "higher is better" if higher_is_better else "lower is better"
#     rows_text = "\n".join(
#         f"{i+1}. {r['router']}: avg={r['avg']:.1f} {param_unit}, min={r['min']:.1f}, max={r['max']:.1f}, spread={r['range']:.1f}"
#         for i, r in enumerate(stats_rows)
#     )
#     return textwrap.dedent(
#         f"""
#         You are writing a professional Wi-Fi comparative report. Do NOT add preamble or conclusions.
#         Write ONLY exactly 5 bullet points. Each bullet must start with •

#         Metric: {param_pretty} ({param_unit})
#         Configuration: {config_label}
#         Floor: {floor_name}
#         Band: {band}
#         Interpretation: {direction_word}

#         Router statistics, best to worst:
#         {rows_text}

#         Rules:
#         - Use only the provided statistics. Do not infer, estimate, or invent rankings, causes, or values.
#         - Output exactly 5 bullets, and nothing else.
#         - Every bullet must begin with •
#         - Bullet 1: summarize the overall pattern and degree of variability across all routers.
#         - Bullet 2: identify the best performer as the first router in the list, and report its exact max and average value.
#         - Bullet 3: identify the mid-tier router or routers only if there are 3 or more total routers; include each name and exact max and  average value.
#         - Bullet 4: identify the worst performer as the last router in the list, and report its exact max and  average value plus the exact gap from the best performer.
#         - Bullet 5: explain the likely technical basis for the separation using radio-frequency or network-performance principles, and keep the explanation scientifically grounded.
#         - Do not add any preamble, closing summary, headings, or extra commentary outside the 5 bullets.
#         - Preserve router names and numeric values exactly as provided.
#         """
#     ).strip()


# def discover_rvr_assets(root: Path, metric_folder: str) -> Dict[str, Any]:
#     metric_dir = root / metric_folder
#     plots_actual: Dict[Tuple[str, str], Path] = {}
#     table_path = None
#     tables_dir = metric_dir / "tables"
#     if tables_dir.exists():
#         candidates = sorted(tables_dir.glob("*_curve_tables.csv"))
#         if candidates:
#             table_path = candidates[0]
#     plot_dir = metric_dir / "plots_actual" / "COMPARE_ROUTERS_BY_FLOOR"
#     if plot_dir.exists():
#         for png in sorted(plot_dir.glob("*.png")):
#             m = re.match(r"^.+?__act__(?P<band>.+?)__floor_(?P<floor>.+?)\.png$", png.name)
#             if m:
#                 plots_actual[(m.group("band"), m.group("floor"))] = png
#     return {"mode": "rvr", "metric_dir": metric_dir, "curve_table_path": table_path, "plots_actual": plots_actual}


# def discover_mesh_assets(root: Path, metric_folder: str) -> Dict[str, Any]:
#     metric_dir = root / metric_folder
#     plot_map: Dict[Tuple[str, str, str], Path] = {}
#     table_path = None
#     tables_dir = metric_dir / "tables"
#     if tables_dir.exists():
#         candidates = sorted(tables_dir.glob("*_mesh_curve_tables.csv"))
#         if candidates:
#             table_path = candidates[0]
#     for png in sorted(metric_dir.rglob("*.png")):
#         if png.parent == metric_dir:
#             continue
#         router = clean_router_name(png.parent.parent.name) if png.parent.parent != metric_dir else ""
#         floor_name = normalize_floor_name(png.parent.name)
#         stem = png.stem
#         if "_" not in stem:
#             continue
#         band, _ = stem.split("_", 1)
#         plot_map[(router, floor_name, normalize_band_value(band))] = png
#     return {"mode": "mesh_compare", "metric_dir": metric_dir, "curve_table_path": table_path, "plots_actual": plot_map}


# def _find_rvr_plot(plot_map: Dict[Tuple[str, str], Path], band: str, floor_name: str) -> Optional[Path]:
#     band = normalize_band_value(band)
#     floor_norm = _safe_slug(normalize_floor_name(floor_name))
#     key = (band, floor_norm)
#     if key in plot_map:
#         return plot_map[key]
#     for (bs, fs), path in plot_map.items():
#         if band in bs and floor_norm.lower() in fs.lower():
#             return path
#     return None


# def _find_mesh_plot(plot_map: Dict[Tuple[str, str, str], Path], router_key: str, floor_name: str, band: str) -> Optional[Path]:
#     key = (clean_router_name(router_key), normalize_floor_name(floor_name), normalize_band_value(band))
#     if key in plot_map:
#         return plot_map[key]
#     return None


# def _compute_stats(subset: pd.DataFrame, higher_is_better: bool) -> List[Dict[str, Any]]:
#     y_col = "p50" if "p50" in subset.columns else ("mean" if "mean" in subset.columns else None)
#     if not y_col:
#         return []
#     rows: List[Dict[str, Any]] = []
#     for router_key in subset["router_key"].astype(str).unique():
#         part = subset[subset["router_key"].astype(str) == router_key]
#         s = pd.to_numeric(part[y_col], errors="coerce").dropna()
#         if s.empty:
#             continue
#         display = clean_router_name(str(part.get("router_display", pd.Series([router_key])).iloc[0]))
#         rows.append(
#             {
#                 "router": display,
#                 "router_key": clean_router_name(router_key),
#                 "avg": float(s.mean()),
#                 "min": float(s.min()),
#                 "max": float(s.max()),
#                 "range": float(s.max() - s.min()),
#             }
#         )
#     rows.sort(key=lambda r: r["avg"], reverse=higher_is_better)
#     return rows


# def _set_landscape(section) -> None:
#     section.orientation = WD_ORIENT.LANDSCAPE
#     section.page_width, section.page_height = section.page_height, section.page_width
#     section.left_margin = Inches(0.6)
#     section.right_margin = Inches(0.6)
#     section.top_margin = Inches(0.7)
#     section.bottom_margin = Inches(0.7)


# def _set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
#     tc = cell._tc
#     tcPr = tc.get_or_add_tcPr()
#     tcMar = tcPr.first_child_found_in("w:tcMar")
#     if tcMar is None:
#         tcMar = OxmlElement("w:tcMar")
#         tcPr.append(tcMar)
#     for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
#         element = tcMar.find(qn(f"w:{tag}"))
#         if element is None:
#             element = OxmlElement(f"w:{tag}")
#             tcMar.append(element)
#         element.set(qn("w:w"), str(value))
#         element.set(qn("w:type"), "dxa")


# def _add_heading(doc: Document, text: str, level: int = 1) -> None:
#     p = doc.add_heading(level=level)
#     run = p.add_run(text)
#     run.bold = True


# def _add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
#     p = doc.add_paragraph()
#     run = p.add_run(text)
#     run.italic = italic
#     run.font.size = Pt(10.5)


# def _add_plot(doc: Document, plot_path: Path) -> None:
#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     p.add_run().add_picture(str(plot_path), width=Inches(9.2))


# def _add_router_cards(doc: Document, router_cards: List[Dict[str, Any]]) -> None:
#     if not router_cards:
#         return
#     cols = 2
#     rows = math.ceil(len(router_cards) / cols)
#     table = doc.add_table(rows=rows, cols=cols)
#     table.alignment = WD_TABLE_ALIGNMENT.CENTER
#     table.style = "Table Grid"
#     table.autofit = False
#     card_width = 4.85
#     heatmap_width = 4.35
#     scale_width = 4.0

#     for idx, card in enumerate(router_cards):
#         cell = table.cell(idx // cols, idx % cols)
#         cell.width = Inches(card_width)
#         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#         _set_cell_margins(cell, top=90, start=90, bottom=90, end=90)

#         p = cell.paragraphs[0]
#         p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         name_run = p.add_run(card["display"])
#         name_run.bold = True
#         name_run.font.size = Pt(10)

#         if card.get("scenario_label"):
#             label_p = cell.add_paragraph()
#             label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             label_run = label_p.add_run(card["scenario_label"])
#             label_run.italic = True
#             label_run.font.size = Pt(9)

#         if card.get("heatmap") and Path(card["heatmap"]).exists():
#             hp = cell.add_paragraph()
#             hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             hp.add_run().add_picture(str(card["heatmap"]), width=Inches(heatmap_width))
#         else:
#             miss = cell.add_paragraph()
#             miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             miss.add_run("Exact heatmap match not found").italic = True

#         if card.get("scale") and Path(card["scale"]).exists():
#             sp = cell.add_paragraph()
#             sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             sp.add_run().add_picture(str(card["scale"]), width=Inches(scale_width))
#         else:
#             miss = cell.add_paragraph()
#             miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             miss.add_run("Matching color scale not found").italic = True

#     total_cells = rows * cols
#     for idx in range(len(router_cards), total_cells):
#         table.cell(idx // cols, idx % cols).text = ""


# def _add_bullets(doc: Document, bullets: List[str]) -> None:
#     for bullet in bullets:
#         p = doc.add_paragraph(style="List Bullet")
#         run = p.add_run(bullet)
#         run.font.size = Pt(10.5)


# def _collect_rvr_router_cards(registry: AssetRegistry, subset: pd.DataFrame, parameter_key: str, floor_name: str, band: str) -> List[Dict[str, Any]]:
#     cards: List[Dict[str, Any]] = []
#     for router_key in subset["router_key"].astype(str).unique():
#         part = subset[subset["router_key"].astype(str) == router_key]
#         display = clean_router_name(str(part["router_display"].iloc[0] if "router_display" in part.columns else router_key))
#         heatmap, scale = registry.get_pair(router_key, parameter_key, floor_name, band)
#         cards.append({"router_key": clean_router_name(router_key), "display": display, "heatmap": heatmap, "scale": scale})
#     return cards


# def _collect_mesh_cards(with_registry: AssetRegistry, without_registry: AssetRegistry, router_key: str, parameter_key: str, floor_name: str, band: str) -> List[Dict[str, Any]]:
#     display = clean_router_name(router_key)
#     with_heatmap, with_scale = with_registry.get_pair(router_key, parameter_key, floor_name, band)
#     wo_heatmap, wo_scale = without_registry.get_pair(router_key, parameter_key, floor_name, band)
#     return [
#         {"router_key": display, "display": display, "scenario_label": "With mesh", "heatmap": with_heatmap, "scale": with_scale},
#         {"router_key": display, "display": display, "scenario_label": "Without mesh", "heatmap": wo_heatmap, "scale": wo_scale},
#     ]


# def _append_audit_rows(audit_rows: List[Dict[str, Any]], cards: List[Dict[str, Any]], parameter_key: str, floor_name: str, band: str, section_type: str):
#     for card in cards:
#         heatmap = card.get("heatmap")
#         scale = card.get("scale")
#         status = "exact_pair"
#         if not heatmap and not scale:
#             status = "missing_heatmap_and_scale"
#         elif not heatmap:
#             status = "missing_heatmap"
#         elif not scale:
#             status = "missing_scale"
#         audit_rows.append(
#             {
#                 "section_type": section_type,
#                 "router_key": card.get("router_key", ""),
#                 "display": card.get("display", ""),
#                 "scenario_label": card.get("scenario_label", ""),
#                 "parameter_key": canonical_metric_key(parameter_key) or parameter_key,
#                 "floor_name": normalize_floor_name(floor_name),
#                 "band": normalize_band_value(band),
#                 "heatmap_path": str(heatmap) if heatmap else "",
#                 "scale_path": str(scale) if scale else "",
#                 "status": status,
#             }
#         )


# def _write_asset_audit(output_path: Path, audit_rows: List[Dict[str, Any]]) -> Optional[Path]:
#     if not audit_rows:
#         return None
#     audit_path = output_path.with_name(f"{output_path.stem}_asset_audit.csv")
#     with audit_path.open("w", encoding="utf-8", newline="") as handle:
#         writer = csv.DictWriter(handle, fieldnames=list(audit_rows[0].keys()))
#         writer.writeheader()
#         writer.writerows(audit_rows)
#     return audit_path


# def generate_report(
#     rvr_outputs_root: Path,
#     extracted_root: Path,
#     output_path: Path,
#     metric_folders: List[str],
#     config_label: str = "Standard",
#     ai_model: str = "gemma3:4b",
#     ai_base_url: str = "http://localhost:11434",
#     use_ai: bool = True,
#     progress_cb: Optional[Callable[[int, int, str], None]] = None,
#     mode: str = "rvr",
#     compare_outputs_root: Optional[Path] = None,
#     extracted_roots_by_scenario: Optional[Dict[str, Path]] = None,
#     csv_outputs_root: Optional[Path] = None,
# ) -> Path:
#     doc = Document()
#     _set_landscape(doc.sections[0])
#     title = doc.add_paragraph()
#     title.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     run = title.add_run("Wi-Fi Comparative Analysis Report")
#     run.bold = True
#     run.font.size = Pt(18)
#     subtitle = "Mesh vs No Mesh" if mode == "mesh_compare" else "Router-to-router comparison"
#     _add_paragraph(doc, subtitle, italic=True)

#     total = max(1, len(metric_folders))
#     registry = AssetRegistry.from_roots(extracted_root=extracted_root, csv_outputs_root=csv_outputs_root)
#     with_registry = without_registry = AssetRegistry([])
#     if mode == "mesh_compare" and extracted_roots_by_scenario:
#         with_registry = AssetRegistry.from_roots(extracted_root=extracted_roots_by_scenario.get("with_mesh", Path(".")))
#         without_registry = AssetRegistry.from_roots(extracted_root=extracted_roots_by_scenario.get("without_mesh", Path(".")))

#     audit_rows: List[Dict[str, Any]] = []

#     for idx, metric_folder in enumerate(metric_folders, start=1):
#         metric_folder = canonical_metric_key(metric_folder) or metric_folder
#         if progress_cb:
#             progress_cb(idx, total, f"Preparing {metric_folder}")
#         param_pretty, param_unit, higher_is_better = _get_param_info(metric_folder)
#         heading_suffix = f" — {config_label}" if config_label and config_label.strip().lower() not in {"", "standard"} else ""
#         _add_heading(doc, f"Range vs {param_pretty}{heading_suffix}", level=1)

#         if mode == "mesh_compare":
#             if compare_outputs_root is None:
#                 raise ValueError("compare_outputs_root is required for mesh_compare mode")
#             assets = discover_mesh_assets(compare_outputs_root, metric_folder)
#         else:
#             assets = discover_rvr_assets(rvr_outputs_root, metric_folder)

#         curve_table_path = assets.get("curve_table_path")
#         if not curve_table_path or not Path(curve_table_path).exists():
#             _add_paragraph(doc, f"No curve table was found for {param_pretty}.")
#             continue

#         curve_table = pd.read_csv(curve_table_path)
#         if curve_table.empty:
#             _add_paragraph(doc, f"The curve table for {param_pretty} is empty.")
#             continue
#         if "band" in curve_table.columns:
#             curve_table["band"] = curve_table["band"].astype(str).map(normalize_band_value)
#         if "floor_name" in curve_table.columns:
#             curve_table["floor_name"] = curve_table["floor_name"].astype(str).map(normalize_floor_name)
#         if "router_key" in curve_table.columns:
#             curve_table["router_key"] = curve_table["router_key"].astype(str).map(clean_router_name)
#         if "router_display" in curve_table.columns:
#             curve_table["router_display"] = curve_table["router_display"].astype(str).map(clean_router_name)

#         intro = f"This section compares {param_pretty} across the available floors and bands." if mode == "rvr" else f"This section compares with-mesh and without-mesh behavior for {param_pretty}."
#         _add_paragraph(doc, intro, italic=True)

#         if mode == "mesh_compare":
#             dimensions = (
#                 curve_table[["router_key", "floor_name", "band"]]
#                 .drop_duplicates()
#                 .sort_values(["router_key", "floor_name", "band"])
#                 .to_dict("records")
#             )
#             for row in dimensions:
#                 router_key = clean_router_name(row["router_key"])
#                 floor_name = normalize_floor_name(str(row["floor_name"]))
#                 band = normalize_band_value(str(row["band"]))
#                 subset = curve_table[
#                     (curve_table["router_key"].astype(str) == router_key)
#                     & (curve_table["floor_name"].astype(str) == floor_name)
#                     & (curve_table["band"].astype(str) == band)
#                 ]
#                 _add_heading(doc, f"{router_key} — {floor_name} — {band}", level=2)
#                 plot_path = _find_mesh_plot(assets["plots_actual"], router_key, floor_name, band)
#                 if plot_path:
#                     _add_plot(doc, plot_path)
#                 cards = _collect_mesh_cards(with_registry, without_registry, router_key, metric_folder, floor_name, band)
#                 _append_audit_rows(audit_rows, cards, metric_folder, floor_name, band, section_type="mesh_compare")
#                 _add_router_cards(doc, cards)
#                 stats_rows = []
#                 for scenario in ["with_mesh", "without_mesh"]:
#                     part = subset[subset["scenario"].astype(str) == scenario]
#                     if part.empty:
#                         continue
#                     y_col = "p50" if "p50" in part.columns else ("mean" if "mean" in part.columns else None)
#                     if not y_col:
#                         continue
#                     vals = pd.to_numeric(part[y_col], errors="coerce").dropna()
#                     if vals.empty:
#                         continue
#                     stats_rows.append(
#                         {
#                             "router": part["scenario_label"].iloc[0] if "scenario_label" in part.columns else scenario,
#                             "avg": float(vals.mean()),
#                             "min": float(vals.min()),
#                             "max": float(vals.max()),
#                             "range": float(vals.max() - vals.min()),
#                         }
#                     )
#                 stats_rows.sort(key=lambda r: r["avg"], reverse=higher_is_better)
#                 bullets = _fallback_bullets(stats_rows, param_pretty, param_unit, floor_name, band, higher_is_better)
#                 _add_bullets(doc, bullets)
#                 doc.add_paragraph()
#         else:
#             bands = sorted(curve_table["band"].astype(str).unique())
#             floors = sorted(curve_table["floor_name"].astype(str).unique())
#             for band in bands:
#                 for floor_name in floors:
#                     subset = curve_table[
#                         (curve_table["band"].astype(str) == band) & (curve_table["floor_name"].astype(str) == floor_name)
#                     ]
#                     if subset.empty:
#                         continue
#                     _add_heading(doc, f"{floor_name} — {band}", level=2)
#                     plot_path = _find_rvr_plot(assets["plots_actual"], band, floor_name)
#                     if plot_path:
#                         _add_plot(doc, plot_path)
#                     cards = _collect_rvr_router_cards(registry, subset, metric_folder, floor_name, band)
#                     _append_audit_rows(audit_rows, cards, metric_folder, floor_name, band, section_type="rvr")
#                     _add_router_cards(doc, cards)
#                     stats_rows = _compute_stats(subset, higher_is_better)
#                     bullets = []
#                     if use_ai and stats_rows:
#                         ai_text = call_ollama(_analysis_prompt(param_pretty, param_unit, floor_name, band, config_label, stats_rows, higher_is_better), model=ai_model, base_url=ai_base_url)
#                         bullets = _parse_bullets(ai_text)
#                     if not bullets:
#                         bullets = _fallback_bullets(stats_rows, param_pretty, param_unit, floor_name, band, higher_is_better)
#                     _add_bullets(doc, bullets)
#                     doc.add_paragraph()
#         if idx < total:
#             doc.add_page_break()

#     doc.save(str(output_path))
#     _write_asset_audit(output_path, audit_rows)
#     return output_path


# def streamlit_report_card(current_router_dir: Optional[Path], rvr_outputs_root: Optional[Path], extracted_root: Optional[Path], step_label: str = "Step 6") -> None:
#     import streamlit as st

#     st.markdown('<div class="card">', unsafe_allow_html=True)
#     st.markdown(
#         f"""
#         <div class="card-title">
#           <h2><span class="step">{step_label}</span> Professional DOCX Report</h2>
#         </div>
#         <div class="subtle">Graph → exact matched heatmaps + exact color scales → analysis bullets. A companion asset audit CSV is written with every report.</div>
#         """,
#         unsafe_allow_html=True,
#     )

#     if current_router_dir is None:
#         st.info("Choose or load a router folder first.")
#         st.markdown("</div>", unsafe_allow_html=True)
#         return

#     mode_label = st.radio("Report mode", ["Parameter vs Range", "Mesh vs No Mesh"], horizontal=True, key="report_mode")
#     mode = "mesh_compare" if mode_label == "Mesh vs No Mesh" else "rvr"

#     metric_source_root = current_router_dir / ("compare_outputs" if mode == "mesh_compare" else "rvr_outputs")
#     metric_dirs = sorted([p.name for p in metric_source_root.iterdir() if p.is_dir()]) if metric_source_root.exists() else []
#     selected_metrics = st.multiselect("Metrics", metric_dirs, default=metric_dirs[:1], key=f"report_metrics_{mode}")
#     config_label = st.text_input("Configuration label", value=("Mesh comparison" if mode == "mesh_compare" else "With Mesh"), key=f"report_cfg_{mode}")
#     use_ai = st.checkbox("Use local AI analysis (Ollama)", value=False, key=f"report_use_ai_{mode}")
#     ai_model = st.text_input("AI model", value="gemma3:4b", key=f"report_ai_model_{mode}")

#     run_button = st.button("Generate DOCX report", key=f"report_generate_{mode}", width="stretch")
#     if run_button:
#         if not selected_metrics:
#             st.error("Select at least one metric.")
#         else:
#             out_path = current_router_dir / ("mesh_compare_report.docx" if mode == "mesh_compare" else "comparative_report.docx")
#             scenario_roots = None
#             if mode == "mesh_compare":
#                 with_root = current_router_dir / "compare_inputs" / "with_mesh_extracted"
#                 without_root = current_router_dir / "compare_inputs" / "without_mesh_extracted"
#                 scenario_roots = {"with_mesh": with_root, "without_mesh": without_root}
#             try:
#                 generate_report(
#                     rvr_outputs_root=rvr_outputs_root or (current_router_dir / "rvr_outputs"),
#                     extracted_root=extracted_root or (current_router_dir / "extracted"),
#                     output_path=out_path,
#                     metric_folders=selected_metrics,
#                     config_label=config_label,
#                     ai_model=ai_model,
#                     use_ai=use_ai,
#                     mode=mode,
#                     compare_outputs_root=(current_router_dir / "compare_outputs"),
#                     extracted_roots_by_scenario=scenario_roots,
#                     csv_outputs_root=(current_router_dir / "csv_outputs"),
#                 )
#                 audit_path = out_path.with_name(f"{out_path.stem}_asset_audit.csv")
#                 st.success(f"Report created: {out_path.name}")
#                 st.download_button("Download report", data=out_path.read_bytes(), file_name=out_path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", width="stretch")
#                 if audit_path.exists():
#                     st.download_button("Download asset audit CSV", data=audit_path.read_bytes(), file_name=audit_path.name, mime="text/csv", width="stretch")
#             except Exception as exc:
#                 st.error(f"Report generation failed: {exc}")

#     st.markdown("</div>", unsafe_allow_html=True)



from __future__ import annotations

import csv
import math
import re
import textwrap
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

import pandas as pd
import requests
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from asset_registry import AssetRegistry
from metadata_utils import PARAM_PRETTY, canonical_metric_key, clean_router_name, normalize_band_value, normalize_floor_name

_PARAM_UNITS = {
    "signal_strength": "dBm",
    "secondary_signal_strength": "dBm",
    "tertiary_signal_strength": "dBm",
    "snr": "dB",
    "noise": "dBm",
    "data_rate": "Mbps",
    "throughput": "Mbps",
    "channel_utilization": "%",
    "channel_interference": "dB",
    "channel_width": "MHz",
    "spectrum_channel_power": "dBm",
    "network_health": "score",
    "network_issues": "count",
    "number_of_access_points": "count",
}

_HIGHER_IS_BETTER = {
    "signal_strength": True,
    "secondary_signal_strength": True,
    "tertiary_signal_strength": True,
    "snr": True,
    "noise": False,
    "data_rate": True,
    "throughput": True,
    "channel_utilization": False,
    "channel_interference": False,
    "channel_width": True,
    "spectrum_channel_power": True,
    "network_health": True,
    "network_issues": False,
    "number_of_access_points": False,
}


def _safe_slug(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]+", "_", str(text)).strip("_")


def _fmt_value(value: float, unit: str) -> str:
    return f"{float(value):.1f} {unit}" if unit else f"{float(value):.1f}"


def _positive_gap(best_avg: float, worst_avg: float, higher_is_better: bool) -> float:
    gap = float(best_avg) - float(worst_avg) if higher_is_better else float(worst_avg) - float(best_avg)
    return abs(gap)


def _router_list_text(rows: List[Dict[str, Any]], unit: str, *, include_max: bool = False) -> str:
    parts: List[str] = []
    for row in rows:
        text = f"{row['router']} ({_fmt_value(row['avg'], unit)} avg"
        if include_max:
            text += f", {_fmt_value(row['max'], unit)} max"
        text += ")"
        parts.append(text)
    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0]
    if len(parts) == 2:
        return f"{parts[0]} and {parts[1]}"
    return ", ".join(parts[:-1]) + f", and {parts[-1]}"


def _metric_specific_reason(param_key: str, band: str, higher_is_better: bool) -> str:
    pkey = canonical_metric_key(param_key) or param_key
    band_text = normalize_band_value(band)

    if pkey in {"signal_strength", "secondary_signal_strength", "tertiary_signal_strength", "spectrum_channel_power"}:
        base = "antenna pattern, receiver sensitivity, spatial-stream behavior, and attenuation handling along the surveyed path"
    elif pkey in {"snr", "noise", "channel_interference", "channel_utilization"}:
        base = "co-channel activity, noise floor control, interference rejection, and airtime management"
    elif pkey in {"throughput", "data_rate"}:
        base = "modulation-and-coding stability, retry behavior, airtime efficiency, and backhaul effectiveness under load"
    elif pkey in {"network_health", "network_issues"}:
        base = "link stability, retransmission behavior, and the way the platform manages contention and roaming events"
    elif pkey in {"number_of_access_points"}:
        base = "cell planning, roaming decisions, and how aggressively the client transitions between available access points"
    else:
        base = "radio-chain performance, propagation loss, and how efficiently the platform handles interference and airtime"

    if band_text.startswith("6"):
        band_note = " At 6 GHz, the separation is usually amplified by the higher free-space loss and weaker wall penetration."
    elif band_text.startswith("5"):
        band_note = " On 5 GHz, the separation often reflects how well each router preserves link quality as attenuation rises with distance."
    elif band_text.startswith("2.4"):
        band_note = " On 2.4 GHz, congestion tolerance and interference handling usually play a larger role than pure path loss alone."
    else:
        band_note = ""

    direction = "higher" if higher_is_better else "lower"
    return f"The separation is most plausibly driven by differences in {base}; for this metric, {direction} values indicate the stronger result.{band_note}"


def _get_param_info(param_key: str) -> Tuple[str, str, bool]:
    pkey = canonical_metric_key(param_key) or param_key
    return (
        PARAM_PRETTY.get(pkey, pkey.replace("_", " ").title()),
        _PARAM_UNITS.get(pkey, ""),
        _HIGHER_IS_BETTER.get(pkey, True),
    )


def call_ollama(prompt: str, model: str = "gemma3:4b", base_url: str = "http://localhost:11434", timeout: int = 90) -> str:
    errors = []

    try:
        resp = requests.post(
            f"{base_url.rstrip('/')}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.2, "top_p": 0.9, "num_predict": 220},
            },
            timeout=timeout,
        )
        if resp.status_code == 404:
            raise requests.exceptions.HTTPError(f"404 at {base_url.rstrip('/')}/api/generate")
        resp.raise_for_status()
        body = resp.json()
        if isinstance(body, dict) and "response" in body and body.get("response"):
            return _strip_preamble(str(body.get("response", "")).strip())
    except Exception as exc:
        errors.append(f"legacy /api/generate: {exc}")

    try:
        resp = requests.post(
            f"{base_url.rstrip('/')}/v1/chat/completions",
            json={
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.2,
                "max_tokens": 220,
            },
            timeout=timeout,
        )
        resp.raise_for_status()
        body = resp.json()
        choices = body.get("choices", [])
        if choices:
            content = choices[0].get("message", {}).get("content")
            if content:
                return _strip_preamble(str(content).strip())
    except Exception as exc:
        errors.append(f"v1/chat/completions: {exc}")

    try:
        resp = requests.post(
            f"{base_url.rstrip('/')}/v1/completions",
            json={
                "model": model,
                "prompt": prompt,
                "temperature": 0.2,
                "max_tokens": 220,
            },
            timeout=timeout,
        )
        resp.raise_for_status()
        body = resp.json()
        choices = body.get("choices", [])
        if choices:
            text = choices[0].get("text") or choices[0].get("message", {}).get("content")
            if text:
                return _strip_preamble(str(text).strip())
    except Exception as exc:
        errors.append(f"v1/completions: {exc}")

    return f"[AI: {' | '.join(errors)}]"


def _strip_preamble(text: str) -> str:
    if not text:
        return text
    text = text.strip()
    preamble_patterns = [
        r"^Here'?s?\s+(?:a\s+)?(?:comparative\s+)?(?:report|summary|analysis)[^:]*:\s*",
        r"^Based on\s+the\s+provided\s+data[^:]*:\s*",
        r"^Here are\s+",
        r"^Here'?s?\s+",
    ]
    for pattern in preamble_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
    conclusion_patterns = [
        r"\n*(?:This|These|The above)\s+(?:analysis|summary|points?|bullets?).*$",
        r"\n*(?:Overall|In summary|In conclusion)[^.]*$",
        r"\n*Feel free to.*$",
    ]
    for pattern in conclusion_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
    return text.strip()


def _parse_bullets(text: str) -> List[str]:
    text = _strip_preamble(text)
    lines: List[str] = []
    for line in (text or "").splitlines():
        line = line.strip()
        if not line:
            continue
        line = re.sub(r"^[\-•*\d\.\)\s]+", "", line)
        line = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", line)
        line = line.strip()
        if len(line) >= 10:
            lines.append(line)
    return lines


def _technical_reason_prompt(param_pretty: str, param_unit: str, floor_name: str, band: str, stats_rows: List[Dict[str, Any]], higher_is_better: bool) -> str:
    direction_word = "higher values are stronger" if higher_is_better else "lower values are stronger"
    rows_text = "\n".join(
        f"- {row['router']}: avg={row['avg']:.1f} {param_unit}, max={row['max']:.1f} {param_unit}, min={row['min']:.1f} {param_unit}"
        for row in stats_rows
    )
    return textwrap.dedent(
        f"""
        Write exactly one bullet line that starts with •.
        Explain the likely technical reason for the separation in a Wi-Fi comparative graph.
        Stay grounded in radio/network science. Do not change the ranking or mention any unsupported facts.

        Metric: {param_pretty} ({param_unit})
        Floor: {floor_name}
        Band: {band}
        Interpretation: {direction_word}

        Router statistics:
        {rows_text}
        """
    ).strip()


def _build_summary_bullets(
    stats_rows: List[Dict[str, Any]],
    param_key: str,
    param_pretty: str,
    param_unit: str,
    floor_name: str,
    band: str,
    higher_is_better: bool,
    *,
    use_ai: bool = False,
    ai_model: str = "gemma3:4b",
    ai_base_url: str = "http://localhost:11434",
) -> List[str]:
    if not stats_rows:
        return [f"No usable data was available for {floor_name} on {band}."]

    best = stats_rows[0]
    worst = stats_rows[-1]
    mid_rows = stats_rows[1:-1] if len(stats_rows) >= 3 else []
    best_to_worst_gap = _positive_gap(best["avg"], worst["avg"], higher_is_better)
    within_router_ranges = [float(row["range"]) for row in stats_rows]
    min_range = min(within_router_ranges)
    max_range = max(within_router_ranges)

    if len(stats_rows) == 1:
        pattern_text = (
            f"The comparative study graph for {floor_name} on {band} contains only one router trace, so it shows a single baseline with "
            f"{_fmt_value(best['range'], param_unit)} of within-curve spread across distance bins."
        )
    else:
        pattern_text = (
            f"The comparative study graph for {floor_name} on {band} shows a clear best-to-worst average separation of "
            f"{_fmt_value(best_to_worst_gap, param_unit)}, while within-router variability spans "
            f"{_fmt_value(min_range, param_unit)} to {_fmt_value(max_range, param_unit)} across the plotted distance bins."
        )

    bullet_1 = pattern_text
    bullet_2 = (
        f"{best['router']} is the best performer with an average of {_fmt_value(best['avg'], param_unit)} "
        f"and a maximum of {_fmt_value(best['max'], param_unit)}."
    )

    if mid_rows:
        bullet_3 = f"The mid-tier router set is {_router_list_text(mid_rows, param_unit, include_max=True)}."
    else:
        bullet_3 = "No mid-tier router classification applies here because fewer than three routers are present in this floor-and-band comparison."

    bullet_4 = (
        f"{worst['router']} is the weakest performer with an average of {_fmt_value(worst['avg'], param_unit)} "
        f"and a maximum of {_fmt_value(worst['max'], param_unit)}; the best-to-worst average gap is {_fmt_value(best_to_worst_gap, param_unit)}."
    )

    bullet_5 = _metric_specific_reason(param_key, band, higher_is_better)
    if use_ai and stats_rows:
        try:
            ai_text = call_ollama(
                _technical_reason_prompt(param_pretty, param_unit, floor_name, band, stats_rows, higher_is_better),
                model=ai_model,
                base_url=ai_base_url,
            )
            ai_lines = _parse_bullets(ai_text)
            if ai_lines:
                bullet_5 = ai_lines[0]
        except Exception:
            pass

    return [bullet_1, bullet_2, bullet_3, bullet_4, bullet_5]


def discover_rvr_assets(root: Path, metric_folder: str) -> Dict[str, Any]:
    metric_dir = root / metric_folder
    plots_actual: Dict[Tuple[str, str], Path] = {}
    table_path = None
    tables_dir = metric_dir / "tables"
    if tables_dir.exists():
        candidates = sorted(tables_dir.glob("*_curve_tables.csv"))
        if candidates:
            table_path = candidates[0]
    plot_dir = metric_dir / "plots_actual" / "COMPARE_ROUTERS_BY_FLOOR"
    if plot_dir.exists():
        for png in sorted(plot_dir.glob("*.png")):
            m = re.match(r"^.+?__act__(?P<band>.+?)__floor_(?P<floor>.+?)\.png$", png.name)
            if m:
                plots_actual[(m.group("band"), m.group("floor"))] = png
    return {"mode": "rvr", "metric_dir": metric_dir, "curve_table_path": table_path, "plots_actual": plots_actual}


def _resolve_metric_dir(root: Path, metric_folder: str) -> Path:
    root = Path(root)
    metric_key = canonical_metric_key(metric_folder) or metric_folder
    direct_child = root / metric_key
    if direct_child.exists() and direct_child.is_dir():
        return direct_child
    root_key = canonical_metric_key(root.name) or root.name
    if root_key == metric_key:
        return root
    return direct_child


def discover_mesh_assets(root: Path, metric_folder: str) -> Dict[str, Any]:
    metric_dir = _resolve_metric_dir(Path(root), metric_folder)
    plot_map: Dict[Tuple[str, str, str], Path] = {}
    table_path = None
    tables_dir = metric_dir / "tables"
    if tables_dir.exists():
        candidates = sorted(tables_dir.glob("*_mesh_curve_tables.csv"))
        if candidates:
            table_path = candidates[0]
    for png in sorted(metric_dir.rglob("*.png")):
        if png.parent == metric_dir:
            continue
        router = clean_router_name(png.parent.parent.name) if png.parent.parent != metric_dir else ""
        floor_name = normalize_floor_name(png.parent.name)
        stem = png.stem
        if "_" not in stem:
            continue
        band, _ = stem.split("_", 1)
        plot_map[(router, floor_name, normalize_band_value(band))] = png
    return {"mode": "mesh_compare", "metric_dir": metric_dir, "curve_table_path": table_path, "plots_actual": plot_map}


def _find_rvr_plot(plot_map: Dict[Tuple[str, str], Path], band: str, floor_name: str) -> Optional[Path]:
    band = normalize_band_value(band)
    floor_norm = _safe_slug(normalize_floor_name(floor_name))
    key = (band, floor_norm)
    if key in plot_map:
        return plot_map[key]
    for (bs, fs), path in plot_map.items():
        if band in bs and floor_norm.lower() in fs.lower():
            return path
    return None


def _find_mesh_plot(plot_map: Dict[Tuple[str, str, str], Path], router_key: str, floor_name: str, band: str) -> Optional[Path]:
    key = (clean_router_name(router_key), normalize_floor_name(floor_name), normalize_band_value(band))
    if key in plot_map:
        return plot_map[key]
    return None


def _pick_y_col(subset: pd.DataFrame) -> Optional[str]:
    for col in ["p50", "mean", "p90", "p10", "max", "min"]:
        if col in subset.columns:
            return col
    return None


def _compute_stats(subset: pd.DataFrame, higher_is_better: bool) -> List[Dict[str, Any]]:
    y_col = _pick_y_col(subset)
    if not y_col:
        return []

    rows: List[Dict[str, Any]] = []
    for router_key in subset["router_key"].astype(str).unique():
        part = subset[subset["router_key"].astype(str) == router_key].copy()
        s = pd.to_numeric(part[y_col], errors="coerce").dropna()
        if s.empty:
            continue

        display_series = part["router_display"] if "router_display" in part.columns else pd.Series([router_key])
        display = clean_router_name(str(display_series.iloc[0]))
        rows.append(
            {
                "router": display,
                "router_key": clean_router_name(router_key),
                "y_col": y_col,
                "avg": float(s.mean()),
                "min": float(s.min()),
                "max": float(s.max()),
                "range": float(s.max() - s.min()),
                "points": int(s.shape[0]),
            }
        )

    rows.sort(
        key=lambda row: (
            row["avg"] if higher_is_better else -row["avg"],
            row["max"] if higher_is_better else -row["max"],
            -row["range"],
            row["router"].lower(),
        ),
        reverse=True,
    )
    return rows


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)


def _set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        element = tcMar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tcMar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.bold = True


def _add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(10.5)


def _add_plot(doc: Document, plot_path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(plot_path), width=Inches(9.2))


def _add_router_cards(doc: Document, router_cards: List[Dict[str, Any]]) -> None:
    if not router_cards:
        return
    cols = 2
    rows = math.ceil(len(router_cards) / cols)
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    card_width = 4.85
    heatmap_width = 4.35
    scale_width = 4.0

    for idx, card in enumerate(router_cards):
        cell = table.cell(idx // cols, idx % cols)
        cell.width = Inches(card_width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_margins(cell, top=90, start=90, bottom=90, end=90)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = p.add_run(card["display"])
        name_run.bold = True
        name_run.font.size = Pt(10)

        if card.get("scenario_label"):
            label_p = cell.add_paragraph()
            label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = label_p.add_run(card["scenario_label"])
            label_run.italic = True
            label_run.font.size = Pt(9)

        if card.get("heatmap") and Path(card["heatmap"]).exists():
            hp = cell.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hp.add_run().add_picture(str(card["heatmap"]), width=Inches(heatmap_width))
        else:
            miss = cell.add_paragraph()
            miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
            miss.add_run("Exact heatmap match not found").italic = True

        if card.get("scale") and Path(card["scale"]).exists():
            sp = cell.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.add_run().add_picture(str(card["scale"]), width=Inches(scale_width))
        else:
            miss = cell.add_paragraph()
            miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
            miss.add_run("Matching color scale not found").italic = True

    total_cells = rows * cols
    for idx in range(len(router_cards), total_cells):
        table.cell(idx // cols, idx % cols).text = ""


def _add_bullets(doc: Document, bullets: List[str]) -> None:
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bullet)
        run.font.size = Pt(10.5)


def _collect_rvr_router_cards(registry: AssetRegistry, subset: pd.DataFrame, parameter_key: str, floor_name: str, band: str) -> List[Dict[str, Any]]:
    cards: List[Dict[str, Any]] = []
    for router_key in subset["router_key"].astype(str).unique():
        part = subset[subset["router_key"].astype(str) == router_key]
        display = clean_router_name(str(part["router_display"].iloc[0] if "router_display" in part.columns else router_key))
        heatmap, scale = registry.get_pair(router_key, parameter_key, floor_name, band)
        cards.append({"router_key": clean_router_name(router_key), "display": display, "heatmap": heatmap, "scale": scale})
    return cards


def _collect_mesh_cards(with_registry: AssetRegistry, without_registry: AssetRegistry, router_key: str, parameter_key: str, floor_name: str, band: str) -> List[Dict[str, Any]]:
    display = clean_router_name(router_key)
    with_heatmap, with_scale = with_registry.get_pair(router_key, parameter_key, floor_name, band)
    wo_heatmap, wo_scale = without_registry.get_pair(router_key, parameter_key, floor_name, band)
    return [
        {"router_key": display, "display": display, "scenario_label": "With mesh", "heatmap": with_heatmap, "scale": with_scale},
        {"router_key": display, "display": display, "scenario_label": "Without mesh", "heatmap": wo_heatmap, "scale": wo_scale},
    ]


def _append_audit_rows(audit_rows: List[Dict[str, Any]], cards: List[Dict[str, Any]], parameter_key: str, floor_name: str, band: str, section_type: str):
    for card in cards:
        heatmap = card.get("heatmap")
        scale = card.get("scale")
        status = "exact_pair"
        if not heatmap and not scale:
            status = "missing_heatmap_and_scale"
        elif not heatmap:
            status = "missing_heatmap"
        elif not scale:
            status = "missing_scale"
        audit_rows.append(
            {
                "section_type": section_type,
                "router_key": card.get("router_key", ""),
                "display": card.get("display", ""),
                "scenario_label": card.get("scenario_label", ""),
                "parameter_key": canonical_metric_key(parameter_key) or parameter_key,
                "floor_name": normalize_floor_name(floor_name),
                "band": normalize_band_value(band),
                "heatmap_path": str(heatmap) if heatmap else "",
                "scale_path": str(scale) if scale else "",
                "status": status,
            }
        )


def _write_asset_audit(output_path: Path, audit_rows: List[Dict[str, Any]]) -> Optional[Path]:
    if not audit_rows:
        return None
    audit_path = output_path.with_name(f"{output_path.stem}_asset_audit.csv")
    with audit_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(audit_rows[0].keys()))
        writer.writeheader()
        writer.writerows(audit_rows)
    return audit_path


def generate_report(
    rvr_outputs_root: Path,
    extracted_root: Path,
    output_path: Path,
    metric_folders: List[str],
    config_label: str = "Standard",
    ai_model: str = "gemma3:4b",
    ai_base_url: str = "http://localhost:11434",
    use_ai: bool = True,
    progress_cb: Optional[Callable[[int, int, str], None]] = None,
    mode: str = "rvr",
    compare_outputs_root: Optional[Path] = None,
    extracted_roots_by_scenario: Optional[Dict[str, Path]] = None,
    csv_outputs_root: Optional[Path] = None,
) -> Path:
    doc = Document()
    _set_landscape(doc.sections[0])
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Wi-Fi Comparative Analysis Report")
    run.bold = True
    run.font.size = Pt(18)
    subtitle = "Mesh vs No Mesh" if mode == "mesh_compare" else "Router-to-router comparison"
    _add_paragraph(doc, subtitle, italic=True)

    total = max(1, len(metric_folders))
    registry = AssetRegistry.from_roots(extracted_root=extracted_root, csv_outputs_root=csv_outputs_root)

    audit_rows: List[Dict[str, Any]] = []

    for idx, metric_folder in enumerate(metric_folders, start=1):
        metric_folder = canonical_metric_key(metric_folder) or metric_folder
        if progress_cb:
            progress_cb(idx, total, f"Preparing {metric_folder}")
        param_pretty, param_unit, higher_is_better = _get_param_info(metric_folder)
        heading_suffix = f" — {config_label}" if config_label and config_label.strip().lower() not in {"", "standard"} else ""
        _add_heading(doc, f"Range vs {param_pretty}{heading_suffix}", level=1)

        metric_compare_dir = None
        if mode == "mesh_compare":
            if compare_outputs_root is None:
                raise ValueError("compare_outputs_root is required for mesh_compare mode")
            assets = discover_mesh_assets(compare_outputs_root, metric_folder)
            metric_compare_dir = assets.get("metric_dir")
        else:
            assets = discover_rvr_assets(rvr_outputs_root, metric_folder)

        curve_table_path = assets.get("curve_table_path")
        if not curve_table_path or not Path(curve_table_path).exists():
            _add_paragraph(doc, f"No curve table was found for {param_pretty}.")
            continue

        curve_table = pd.read_csv(curve_table_path)
        if curve_table.empty:
            _add_paragraph(doc, f"The curve table for {param_pretty} is empty.")
            continue
        if "band" in curve_table.columns:
            curve_table["band"] = curve_table["band"].astype(str).map(normalize_band_value)
        if "floor_name" in curve_table.columns:
            curve_table["floor_name"] = curve_table["floor_name"].astype(str).map(normalize_floor_name)
        if "router_key" in curve_table.columns:
            curve_table["router_key"] = curve_table["router_key"].astype(str).map(clean_router_name)
        if "router_display" in curve_table.columns:
            curve_table["router_display"] = curve_table["router_display"].astype(str).map(clean_router_name)

        intro = f"This section summarizes the comparative-study graph for {param_pretty} using the curve-table statistics." if mode == "rvr" else f"This section compares with-mesh and without-mesh behavior for {param_pretty}."
        _add_paragraph(doc, intro, italic=True)

        if mode == "mesh_compare":
            with_registry = without_registry = AssetRegistry([])
            if metric_compare_dir is not None:
                metric_with_root = metric_compare_dir / "with_mesh_extracted"
                metric_without_root = metric_compare_dir / "without_mesh_extracted"
                if metric_with_root.exists():
                    with_registry = AssetRegistry.from_roots(extracted_root=metric_with_root)
                if metric_without_root.exists():
                    without_registry = AssetRegistry.from_roots(extracted_root=metric_without_root)
            if extracted_roots_by_scenario:
                if not with_registry.records and extracted_roots_by_scenario.get("with_mesh"):
                    with_registry = AssetRegistry.from_roots(extracted_root=extracted_roots_by_scenario.get("with_mesh"))
                if not without_registry.records and extracted_roots_by_scenario.get("without_mesh"):
                    without_registry = AssetRegistry.from_roots(extracted_root=extracted_roots_by_scenario.get("without_mesh"))
            dimensions = (
                curve_table[["router_key", "floor_name", "band"]]
                .drop_duplicates()
                .sort_values(["router_key", "floor_name", "band"])
                .to_dict("records")
            )
            for row in dimensions:
                router_key = clean_router_name(row["router_key"])
                floor_name = normalize_floor_name(str(row["floor_name"]))
                band = normalize_band_value(str(row["band"]))
                subset = curve_table[
                    (curve_table["router_key"].astype(str) == router_key)
                    & (curve_table["floor_name"].astype(str) == floor_name)
                    & (curve_table["band"].astype(str) == band)
                ]
                _add_heading(doc, f"{router_key} — {floor_name} — {band}", level=2)
                plot_path = _find_mesh_plot(assets["plots_actual"], router_key, floor_name, band)
                if plot_path:
                    _add_plot(doc, plot_path)
                cards = _collect_mesh_cards(with_registry, without_registry, router_key, metric_folder, floor_name, band)
                _append_audit_rows(audit_rows, cards, metric_folder, floor_name, band, section_type="mesh_compare")
                _add_router_cards(doc, cards)
                stats_rows = []
                for scenario in ["with_mesh", "without_mesh"]:
                    part = subset[subset["scenario"].astype(str) == scenario]
                    if part.empty:
                        continue
                    y_col = _pick_y_col(part)
                    if not y_col:
                        continue
                    vals = pd.to_numeric(part[y_col], errors="coerce").dropna()
                    if vals.empty:
                        continue
                    stats_rows.append(
                        {
                            "router": part["scenario_label"].iloc[0] if "scenario_label" in part.columns else scenario,
                            "avg": float(vals.mean()),
                            "min": float(vals.min()),
                            "max": float(vals.max()),
                            "range": float(vals.max() - vals.min()),
                        }
                    )
                stats_rows.sort(key=lambda row: row["avg"], reverse=higher_is_better)
                bullets = _build_summary_bullets(
                    stats_rows,
                    metric_folder,
                    param_pretty,
                    param_unit,
                    floor_name,
                    band,
                    higher_is_better,
                    use_ai=use_ai,
                    ai_model=ai_model,
                    ai_base_url=ai_base_url,
                )
                _add_bullets(doc, bullets)
                doc.add_paragraph()
        else:
            bands = sorted(curve_table["band"].astype(str).unique())
            floors = sorted(curve_table["floor_name"].astype(str).unique())
            for band in bands:
                for floor_name in floors:
                    subset = curve_table[
                        (curve_table["band"].astype(str) == band) & (curve_table["floor_name"].astype(str) == floor_name)
                    ]
                    if subset.empty:
                        continue
                    _add_heading(doc, f"{floor_name} — {band}", level=2)
                    plot_path = _find_rvr_plot(assets["plots_actual"], band, floor_name)
                    if plot_path:
                        _add_plot(doc, plot_path)
                    cards = _collect_rvr_router_cards(registry, subset, metric_folder, floor_name, band)
                    _append_audit_rows(audit_rows, cards, metric_folder, floor_name, band, section_type="rvr")
                    _add_router_cards(doc, cards)
                    stats_rows = _compute_stats(subset, higher_is_better)
                    bullets = _build_summary_bullets(
                        stats_rows,
                        metric_folder,
                        param_pretty,
                        param_unit,
                        floor_name,
                        band,
                        higher_is_better,
                        use_ai=use_ai,
                        ai_model=ai_model,
                        ai_base_url=ai_base_url,
                    )
                    _add_bullets(doc, bullets)
                    doc.add_paragraph()
        if idx < total:
            doc.add_page_break()

    doc.save(str(output_path))
    _write_asset_audit(output_path, audit_rows)
    return output_path


def streamlit_report_card(current_router_dir: Optional[Path], rvr_outputs_root: Optional[Path], extracted_root: Optional[Path], step_label: str = "Step 6") -> None:
    import streamlit as st

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="card-title">
          <h2><span class="step">{step_label}</span> Professional DOCX Report</h2>
        </div>
        <div class="subtle">Graph → exact matched heatmaps + exact color scales → graph-aligned summary bullets. A companion asset audit CSV is written with every report.</div>
        """,
        unsafe_allow_html=True,
    )

    if current_router_dir is None:
        st.info("Choose or load a router folder first.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    mode_label = st.radio("Report mode", ["Parameter vs Range", "Mesh vs No Mesh"], horizontal=True, key="report_mode")
    mode = "mesh_compare" if mode_label == "Mesh vs No Mesh" else "rvr"

    metric_source_root = current_router_dir / ("compare_outputs" if mode == "mesh_compare" else "rvr_outputs")
    metric_dirs = sorted([p.name for p in metric_source_root.iterdir() if p.is_dir()]) if metric_source_root.exists() else []
    selected_metrics = st.multiselect("Metrics", metric_dirs, default=metric_dirs[:1], key=f"report_metrics_{mode}")
    config_label = st.text_input("Configuration label", value=("Mesh comparison" if mode == "mesh_compare" else "With Mesh"), key=f"report_cfg_{mode}")
    use_ai = st.checkbox("Use local AI analysis (Ollama)", value=False, key=f"report_use_ai_{mode}")
    ai_model = st.text_input("AI model", value="gemma3:4b", key=f"report_ai_model_{mode}")

    run_button = st.button("Generate DOCX report", key=f"report_generate_{mode}", width="stretch")
    if run_button:
        if not selected_metrics:
            st.error("Select at least one metric.")
        else:
            out_path = current_router_dir / ("mesh_compare_report.docx" if mode == "mesh_compare" else "comparative_report.docx")
            scenario_roots = None
            try:
                generate_report(
                    rvr_outputs_root=rvr_outputs_root or (current_router_dir / "rvr_outputs"),
                    extracted_root=extracted_root or (current_router_dir / "extracted"),
                    output_path=out_path,
                    metric_folders=selected_metrics,
                    config_label=config_label,
                    ai_model=ai_model,
                    use_ai=use_ai,
                    mode=mode,
                    compare_outputs_root=(current_router_dir / "compare_outputs"),
                    extracted_roots_by_scenario=scenario_roots,
                    csv_outputs_root=(current_router_dir / "csv_outputs"),
                )
                audit_path = out_path.with_name(f"{out_path.stem}_asset_audit.csv")
                st.success(f"Report created: {out_path.name}")
                st.download_button("Download report", data=out_path.read_bytes(), file_name=out_path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", width="stretch")
                if audit_path.exists():
                    st.download_button("Download asset audit CSV", data=audit_path.read_bytes(), file_name=audit_path.name, mime="text/csv", width="stretch")
            except Exception as exc:
                st.error(f"Report generation failed: {exc}")

    st.markdown("</div>", unsafe_allow_html=True)
