from __future__ import annotations

import json
import math
import re
from collections import Counter
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


HIGHER_IS_BETTER_BY_METRIC: Dict[str, bool] = {
    "Signal Strength": True,
    "Secondary Signal Strength": True,
    "Tertiary Signal Strength": True,
    "SNR": True,
    "Noise": False,
    "Data Rate": True,
    "Throughput": True,
    "Channel Utilization": False,
    "Channel Interference": False,
    "Channel Width": True,
    "Spectrum Channel Power": True,
    "Network Health": True,
    "Network Issues": False,
    "Number of APs": False,
    "Number of Access Points": False,
}


# -------------------------------------------------
# Developer-only local AI defaults
# -------------------------------------------------
DEFAULT_LOCAL_AI_PROVIDER = "ollama"
DEFAULT_LOCAL_AI_MODEL = "gemma3:4b"
DEFAULT_LOCAL_AI_BASE_URL = "http://127.0.0.1:11434"
DEFAULT_LOCAL_AI_API_KEY = ""
DEFAULT_LOCAL_AI_TEMPERATURE = 0.1
DEFAULT_LOCAL_AI_MAX_TOKENS = 1400
DEFAULT_LOCAL_AI_TIMEOUT_SEC: Optional[int] = None
DEFAULT_LOCAL_AI_EXTRA_INSTRUCTIONS = (
    "Write grounded engineering summaries for WiFi site survey graphs. "
    "Use only the supplied metric context, rankings, and trend data. "
    "Do not invent facts, do not add questions, and do not add conversational endings."
)

FORBIDDEN_REPORT_PATTERNS = [
    r"would you like me to .*",
    r"let me know .*",
    r"if you (?:would like|want),? .*",
    r"i can (?:refine|expand|add more detail).*",
    r"perhaps focusing on .*",
    r"adding more detail .*",
    r"potential troubleshooting steps.*",
    r"feel free to ask.*",
    r"do you want me to .*",
]

GRAPH_SECTION_ORDER = [
    "What the graph compares",
    "Main observations",
    "Distance behavior",
    "Engineering reasoning",
    "Cautions",
]
METRIC_SECTION_ORDER = [
    "Strong patterns",
    "Weak patterns",
    "Distance-related findings",
    "Engineering takeaways",
    "Data limits",
]
OVERALL_SECTION_ORDER = [
    "Scope covered",
    "Cross-metric findings",
    "Floors and bands",
    "Router comparisons",
    "Engineering takeaways",
    "Data limits",
]


@dataclass
class LocalAIConfig:
    provider: str = DEFAULT_LOCAL_AI_PROVIDER
    model: str = DEFAULT_LOCAL_AI_MODEL
    base_url: str = DEFAULT_LOCAL_AI_BASE_URL
    api_key: str = DEFAULT_LOCAL_AI_API_KEY
    temperature: float = DEFAULT_LOCAL_AI_TEMPERATURE
    max_tokens: int = DEFAULT_LOCAL_AI_MAX_TOKENS
    timeout_sec: Optional[int] = DEFAULT_LOCAL_AI_TIMEOUT_SEC
    extra_instructions: str = DEFAULT_LOCAL_AI_EXTRA_INSTRUCTIONS


def _clean_base_url(base_url: str) -> str:
    return (base_url or "").strip().rstrip("/")


def _probe_timeout_seconds(timeout_sec: Optional[int]) -> int:
    if timeout_sec is None:
        return 20
    try:
        t = int(timeout_sec)
    except Exception:
        return 20
    if t <= 0:
        return 20
    return min(t, 20)


def _request_timeout_seconds(timeout_sec: Optional[int]) -> Optional[int]:
    if timeout_sec is None:
        return None
    try:
        t = int(timeout_sec)
    except Exception:
        return None
    if t <= 0:
        return None
    return t


def warm_local_ai(cfg: LocalAIConfig) -> None:
    if cfg.provider != "ollama":
        return
    resp = requests.post(
        f"{_clean_base_url(cfg.base_url)}/api/generate",
        json={"model": cfg.model, "prompt": "", "stream": False, "keep_alive": -1},
        timeout=_probe_timeout_seconds(cfg.timeout_sec),
    )
    resp.raise_for_status()


def unload_local_ai(cfg: LocalAIConfig) -> None:
    if cfg.provider != "ollama":
        return
    resp = requests.post(
        f"{_clean_base_url(cfg.base_url)}/api/generate",
        json={"model": cfg.model, "prompt": "", "stream": False, "keep_alive": 0},
        timeout=_probe_timeout_seconds(cfg.timeout_sec),
    )
    resp.raise_for_status()


def probe_local_ai(cfg: LocalAIConfig) -> Tuple[bool, str]:
    try:
        base_url = _clean_base_url(cfg.base_url)
        probe_timeout = _probe_timeout_seconds(cfg.timeout_sec)

        if cfg.provider == "ollama":
            resp = requests.get(f"{base_url}/api/tags", timeout=probe_timeout)
            resp.raise_for_status()
            data = resp.json()
            models = [m.get("name", "") for m in data.get("models", []) if isinstance(m, dict)]
            if cfg.model and cfg.model not in models:
                return False, f"Connected, but model '{cfg.model}' was not found. Available: {', '.join(models[:12]) or 'none'}"
            return True, f"Connected to Ollama. Models: {', '.join(models[:12]) or 'none'}"

        headers = {"Content-Type": "application/json"}
        if cfg.api_key:
            headers["Authorization"] = f"Bearer {cfg.api_key}"
        resp = requests.get(f"{base_url}/models", headers=headers, timeout=probe_timeout)
        if resp.status_code == 404:
            resp = requests.get(f"{base_url}/v1/models", headers=headers, timeout=probe_timeout)
        resp.raise_for_status()
        data = resp.json()
        models = [m.get("id", "") for m in data.get("data", []) if isinstance(m, dict)]
        if cfg.model and cfg.model not in models:
            return False, f"Connected, but model '{cfg.model}' was not found. Available: {', '.join(models[:12]) or 'none'}"
        return True, f"Connected to local OpenAI-compatible API. Models: {', '.join(models[:12]) or 'none'}"
    except Exception as exc:
        return False, f"Connection failed: {type(exc).__name__}: {exc}"


def _safe_float(value: Any) -> Optional[float]:
    try:
        x = float(value)
    except Exception:
        return None
    if not math.isfinite(x):
        return None
    return x


def _round_or_none(value: Any, ndigits: int = 3) -> Optional[float]:
    x = _safe_float(value)
    return None if x is None else round(x, ndigits)


def _truncate_text(text: str, limit: int = 18000) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[truncated]"


def _clean_generated_report_text(text: str) -> str:
    cleaned_lines: List[str] = []
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        lower = line.lower()
        if any(re.search(pattern, lower) for pattern in FORBIDDEN_REPORT_PATTERNS):
            continue
        cleaned_lines.append(raw_line.rstrip())
    cleaned = "\n".join(cleaned_lines).strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned


def _call_ollama(prompt: str, cfg: LocalAIConfig) -> str:
    resp = requests.post(
        f"{_clean_base_url(cfg.base_url)}/api/generate",
        json={
            "model": cfg.model,
            "prompt": prompt,
            "stream": False,
            "keep_alive": -1,
            "options": {"temperature": cfg.temperature, "num_predict": int(cfg.max_tokens)},
        },
        timeout=_request_timeout_seconds(cfg.timeout_sec),
    )
    resp.raise_for_status()
    data = resp.json()
    if not isinstance(data, dict) or not data.get("response"):
        raise RuntimeError(f"Unexpected Ollama response: {data}")
    return str(data["response"]).strip()


def _call_openai_compatible(prompt: str, cfg: LocalAIConfig) -> str:
    base_url = _clean_base_url(cfg.base_url)
    endpoint = f"{base_url}/chat/completions"
    headers = {"Content-Type": "application/json"}
    if cfg.api_key:
        headers["Authorization"] = f"Bearer {cfg.api_key}"
    payload = {
        "model": cfg.model,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a WiFi site survey analyst. "
                    "Write formal report sections only. "
                    "Use only the provided data, and never add offers or questions."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": cfg.temperature,
        "max_tokens": int(cfg.max_tokens),
    }
    resp = requests.post(endpoint, headers=headers, json=payload, timeout=_request_timeout_seconds(cfg.timeout_sec))
    if resp.status_code == 404 and not endpoint.endswith("/v1/chat/completions"):
        endpoint = f"{base_url}/v1/chat/completions"
        resp = requests.post(endpoint, headers=headers, json=payload, timeout=_request_timeout_seconds(cfg.timeout_sec))
    resp.raise_for_status()
    data = resp.json()
    try:
        return str(data["choices"][0]["message"]["content"]).strip()
    except Exception as exc:
        raise RuntimeError(f"Unexpected chat completion response: {data}") from exc


def run_local_ai(prompt: str, cfg: LocalAIConfig) -> str:
    if cfg.provider == "ollama":
        return _call_ollama(prompt, cfg)
    return _call_openai_compatible(prompt, cfg)


# -----------------------------
# DOCX/report helpers
# -----------------------------
def _configure_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 20, RGBColor(20, 37, 63)),
        ("Heading 1", 15, RGBColor(20, 37, 63)),
        ("Heading 2", 12, RGBColor(36, 71, 114)),
        ("Heading 3", 11, RGBColor(36, 71, 114)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def _extract_markdown_sections(markdown_text: str) -> Tuple[str, Dict[str, str]]:
    title = ""
    sections: Dict[str, List[str]] = {}
    current_heading: Optional[str] = None

    for raw_line in (markdown_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            if current_heading:
                sections.setdefault(current_heading, []).append("")
            continue
        if line.startswith("# ") and not title:
            title = line[2:].strip()
            continue
        if line.startswith("## "):
            current_heading = line[3:].strip()
            sections.setdefault(current_heading, [])
            continue
        if line.startswith("### "):
            sub = line[4:].strip()
            if current_heading:
                sections.setdefault(current_heading, []).append(f"**{sub}**")
            continue
        if current_heading:
            sections.setdefault(current_heading, []).append(line)

    return title, {k: "\n".join(v).strip() for k, v in sections.items()}


def _compose_markdown_report(title: str, section_order: List[str], section_map: Dict[str, str]) -> str:
    lines = [f"# {title}"]
    for heading in section_order:
        body = (section_map.get(heading) or "Not enough supported evidence was available for this section.").strip()
        lines.extend(["", f"## {heading}", body])
    return "\n".join(lines).strip() + "\n"


def _markdown_to_docx(markdown_text: str, title: str, out_path: Path, image_paths: Optional[List[Path]] = None) -> Path:
    doc = Document()
    _configure_doc_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(20, 37, 63)

    title_in_text, sections = _extract_markdown_sections(markdown_text)
    report_title = title_in_text or title

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(report_title)
    meta_run.italic = True
    meta_run.font.name = "Arial"
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor(88, 96, 105)

    for heading, body in sections.items():
        doc.add_heading(heading, level=2)
        for raw_line in body.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if re.match(r"^[-*]\s+", line):
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(re.sub(r"^[-*]\s+", "", line))
            elif line.startswith("**") and line.endswith("**") and len(line) > 4:
                para = doc.add_paragraph()
                run = para.add_run(line.strip("*"))
                run.bold = True
            else:
                doc.add_paragraph(line)

    valid_images = [Path(p) for p in (image_paths or []) if Path(p).exists()]
    if valid_images:
        doc.add_heading("Comparative Study Graphs", level=2)
        for img in valid_images:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = caption.add_run(img.name)
            cap_run.bold = True
            cap_run.font.name = "Arial"
            cap_run.font.size = Pt(9.5)
            try:
                doc.add_picture(str(img), width=Inches(6.4))
            except Exception:
                doc.add_paragraph(f"[Could not embed graph image: {img.name}]")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def save_text_artifacts(
    text: str,
    title: str,
    out_dir: Path,
    stem: str,
    metadata: Optional[Dict[str, Any]] = None,
    image_paths: Optional[List[Path]] = None,
) -> Dict[str, str]:
    out_dir.mkdir(parents=True, exist_ok=True)
    md_path = out_dir / f"{stem}.md"
    docx_path = out_dir / f"{stem}.docx"
    json_path = out_dir / f"{stem}.json"

    cleaned_text = _clean_generated_report_text(text)
    md_path.write_text(cleaned_text, encoding="utf-8")
    _markdown_to_docx(cleaned_text, title=title, out_path=docx_path, image_paths=image_paths)

    payload = dict(metadata or {})
    payload["title"] = title
    payload["text"] = cleaned_text
    payload["image_paths"] = [str(Path(p)) for p in (image_paths or [])]
    json_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")

    return {"markdown": str(md_path), "docx": str(docx_path), "json": str(json_path)}


# -----------------------------
# Metric + graph context extraction
# -----------------------------
def _load_curve_table(metric_output_dir: Path) -> Optional[pd.DataFrame]:
    table_dir = metric_output_dir / "tables"
    if not table_dir.exists():
        return None
    matches = sorted(table_dir.glob("*_curve_tables.csv"))
    if not matches:
        return None
    try:
        return pd.read_csv(matches[0])
    except Exception:
        return None


def _summarize_curves(df: Optional[pd.DataFrame], higher_is_better: bool) -> List[Dict[str, Any]]:
    if df is None or df.empty:
        return []

    required = {"router_key", "floor_name", "band", "dist_ft_mid"}
    if not required.issubset(set(df.columns)):
        return []

    y_col = None
    for candidate in ["p50", "mean", "metric_p50", "metric_mean", "value"]:
        if candidate in df.columns:
            y_col = candidate
            break
    if not y_col:
        return []

    work = df.copy()
    work[y_col] = pd.to_numeric(work[y_col], errors="coerce")
    work["dist_ft_mid"] = pd.to_numeric(work["dist_ft_mid"], errors="coerce")
    work = work.dropna(subset=[y_col, "dist_ft_mid"])
    if work.empty:
        return []

    rows: List[Dict[str, Any]] = []
    for (router_key, floor_name, band), grp in work.groupby(["router_key", "floor_name", "band"], dropna=False):
        grp = grp.sort_values("dist_ft_mid")
        first_val = _safe_float(grp.iloc[0][y_col])
        last_val = _safe_float(grp.iloc[-1][y_col])
        max_dist = _safe_float(grp["dist_ft_mid"].max())
        avg_val = _safe_float(grp[y_col].mean())
        delta = None
        if first_val is not None and last_val is not None:
            delta = last_val - first_val

        trend = "flat"
        if delta is not None:
            if higher_is_better:
                if delta <= -1e-9:
                    trend = "declining_with_distance"
                elif delta >= 1e-9:
                    trend = "improving_with_distance"
            else:
                if delta >= 1e-9:
                    trend = "worsening_with_distance"
                elif delta <= -1e-9:
                    trend = "improving_with_distance"

        rows.append(
            {
                "router_key": str(router_key),
                "floor_name": str(floor_name),
                "band": str(band),
                "bins": int(len(grp)),
                "max_distance_ft": _round_or_none(max_dist, 2),
                "avg_curve_value": _round_or_none(avg_val, 3),
                "start_value": _round_or_none(first_val, 3),
                "end_value": _round_or_none(last_val, 3),
                "end_minus_start": _round_or_none(delta, 3),
                "trend": trend,
            }
        )

    sort_desc = higher_is_better
    rows.sort(
        key=lambda r: (
            r.get("avg_curve_value") is not None,
            r.get("avg_curve_value") if r.get("avg_curve_value") is not None else (-1e18 if sort_desc else 1e18),
        ),
        reverse=sort_desc,
    )
    return rows[:100]


def _load_rankings(metric_output_dir: Path) -> List[Dict[str, Any]]:
    rank_dir = metric_output_dir / "plots_actual" / "COMPARE_ROUTERS_BY_FLOOR"
    if not rank_dir.exists():
        return []

    items: List[Dict[str, Any]] = []
    for path in sorted(rank_dir.glob("*__avg_actual.csv")):
        try:
            df = pd.read_csv(path)
        except Exception:
            continue
        if df.empty:
            continue

        cols = {c.lower(): c for c in df.columns}
        floor_col = cols.get("floor")
        band_col = cols.get("band")
        router_col = cols.get("router_name") or cols.get("router_key")
        rank_col = cols.get("rank_best_to_worst")
        avg_col = cols.get("avg_actual")
        if not all([floor_col, band_col, router_col, rank_col, avg_col]):
            continue

        df = df.sort_values(rank_col)
        ordered = []
        for _, row in df.head(12).iterrows():
            ordered.append(
                {
                    "router": str(row.get(router_col, "")),
                    "rank": int(row.get(rank_col)) if pd.notna(row.get(rank_col)) else None,
                    "avg_actual": _round_or_none(row.get(avg_col), 3),
                }
            )
        if ordered:
            items.append(
                {
                    "file": path.name,
                    "floor": str(df.iloc[0].get(floor_col, "")),
                    "band": str(df.iloc[0].get(band_col, "")),
                    "ordered_best_to_worst": ordered,
                }
            )
    return items[:120]


def build_metric_context(metric_output_dir: Path, metric_display: str) -> Dict[str, Any]:
    higher_is_better = HIGHER_IS_BETTER_BY_METRIC.get(metric_display, True)
    curve_df = _load_curve_table(metric_output_dir)
    curve_summaries = _summarize_curves(curve_df, higher_is_better=higher_is_better)
    rankings = _load_rankings(metric_output_dir)
    plot_paths = sorted(metric_output_dir.rglob("*.png"))

    context: Dict[str, Any] = {
        "metric_display": metric_display,
        "higher_is_better": higher_is_better,
        "available_plot_count": len(plot_paths),
        "plot_examples": [str(p.name) for p in plot_paths[:24]],
        "curve_summary_count": len(curve_summaries),
        "curve_summaries": curve_summaries,
        "ranking_count": len(rankings),
        "rankings": rankings,
    }
    if curve_df is not None and not curve_df.empty:
        context["curve_table_columns"] = [str(c) for c in curve_df.columns]
        context["curve_table_rows"] = int(len(curve_df))
    else:
        context["curve_table_columns"] = []
        context["curve_table_rows"] = 0
    return context


def _load_plot_meta(plot_path: Path) -> Dict[str, Any]:
    meta_path = plot_path.with_suffix(".meta.json")
    if not meta_path.exists():
        return {}
    try:
        return json.loads(meta_path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def infer_plot_info(plot_path: Path) -> Dict[str, Any]:
    info: Dict[str, Any] = {
        "plot_name": plot_path.name,
        "plot_path": str(plot_path),
        "view_mode": "actual" if "plots_actual" in str(plot_path) else ("percent" if "plots_percent" in str(plot_path) else "unknown"),
        "comparison_type": "unknown",
        "band": "",
        "focus_floor": "",
        "focus_router": "",
    }
    lower = plot_path.as_posix().lower()
    if "compare_routers_by_floor" in lower:
        info["comparison_type"] = "compare_routers_by_floor"
    elif "compare_floors_by_router" in lower:
        info["comparison_type"] = "compare_floors_by_router"

    band_match = re.search(r"__(2\.4ghz|5ghz)__", plot_path.name, flags=re.IGNORECASE)
    if band_match:
        info["band"] = band_match.group(1)
    floor_match = re.search(r"__floor_(.+?)\.png$", plot_path.name, flags=re.IGNORECASE)
    router_match = re.search(r"__router_(.+?)\.png$", plot_path.name, flags=re.IGNORECASE)
    if floor_match:
        info["focus_floor"] = floor_match.group(1).replace("_", " ")
    if router_match:
        info["focus_router"] = router_match.group(1).replace("_", " ")
    info["meta"] = _load_plot_meta(plot_path)
    return info


def _norm_text(text: str) -> str:
    text = (text or "").strip().lower()
    text = text.replace("_", " ").replace("-", " ")
    text = re.sub(r"\s+", " ", text)
    return text


def _filter_relevant_curves(metric_context: Dict[str, Any], plot_info: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows = metric_context.get("curve_summaries", []) or []
    band = _norm_text(plot_info.get("band", ""))
    focus_floor = _norm_text(plot_info.get("focus_floor", ""))
    focus_router = _norm_text(plot_info.get("focus_router", ""))
    comparison_type = plot_info.get("comparison_type", "")

    out: List[Dict[str, Any]] = []
    for row in rows:
        row_band = _norm_text(str(row.get("band", "")))
        row_floor = _norm_text(str(row.get("floor_name", "")))
        row_router = _norm_text(str(row.get("router_key", "")))
        if band and row_band != band:
            continue
        if comparison_type == "compare_routers_by_floor" and focus_floor and focus_floor not in row_floor:
            continue
        if comparison_type == "compare_floors_by_router" and focus_router and focus_router not in row_router:
            continue
        out.append(row)
    return out[:24]


def _filter_relevant_rankings(metric_context: Dict[str, Any], plot_info: Dict[str, Any]) -> List[Dict[str, Any]]:
    items = metric_context.get("rankings", []) or []
    band = _norm_text(plot_info.get("band", ""))
    focus_floor = _norm_text(plot_info.get("focus_floor", ""))
    out: List[Dict[str, Any]] = []
    for row in items:
        row_band = _norm_text(str(row.get("band", "")))
        row_floor = _norm_text(str(row.get("floor", "")))
        if band and row_band != band:
            continue
        if focus_floor and focus_floor not in row_floor:
            continue
        out.append(row)
    return out[:10]


def _trend_sentence(trend: str, higher_is_better: bool) -> str:
    mapping = {
        "declining_with_distance": "the metric generally degrades as distance increases",
        "worsening_with_distance": "the metric generally worsens as distance increases",
        "improving_with_distance": "the metric generally improves as distance increases",
        "flat": "the metric is comparatively flat over distance",
    }
    if trend == "improving_with_distance" and higher_is_better:
        return mapping[trend]
    if trend == "improving_with_distance" and not higher_is_better:
        return "the metric generally improves as distance increases"
    return mapping.get(trend, "distance behavior could not be classified")


def _distance_behavior_from_curves(curves: List[Dict[str, Any]], higher_is_better: bool) -> str:
    if not curves:
        return (
            "The available numeric context was not sufficient to establish a reliable distance-behavior statement "
            "for this graph."
        )

    counts = Counter(str(row.get("trend", "flat")) for row in curves)
    dominant = counts.most_common(1)[0][0]
    lines = [
        f"- Relevant curves reviewed: {len(curves)}.",
        f"- Dominant distance pattern: {_trend_sentence(dominant, higher_is_better)}.",
    ]
    if counts.get("declining_with_distance"):
        lines.append(f"- Declining curves: {counts['declining_with_distance']}.")
    if counts.get("worsening_with_distance"):
        lines.append(f"- Worsening curves: {counts['worsening_with_distance']}.")
    if counts.get("improving_with_distance"):
        lines.append(f"- Improving curves: {counts['improving_with_distance']}.")
    if counts.get("flat"):
        lines.append(f"- Flat curves: {counts['flat']}.")

    example_rows = curves[:3]
    for row in example_rows:
        lines.append(
            "- {router} | {floor} | {band}: start {start}, end {end}, delta {delta}.".format(
                router=row.get("router_key", "—"),
                floor=row.get("floor_name", "—"),
                band=row.get("band", "—"),
                start=row.get("start_value", "—"),
                end=row.get("end_value", "—"),
                delta=row.get("end_minus_start", "—"),
            )
        )
    return "\n".join(lines)


def _graph_scope_text(plot_info: Dict[str, Any]) -> str:
    ctype = plot_info.get("comparison_type", "unknown")
    floor = plot_info.get("focus_floor", "") or "the relevant floor"
    router = plot_info.get("focus_router", "") or "the relevant router"
    band = plot_info.get("band", "") or "the relevant band"
    mode = plot_info.get("view_mode", "unknown")

    if ctype == "compare_routers_by_floor":
        return f"This graph compares multiple routers on {floor} for the {band} band using {mode} values."
    if ctype == "compare_floors_by_router":
        return f"This graph compares multiple floors for router {router} on the {band} band using {mode} values."
    return f"This graph summarizes {band} comparative behavior using {mode} values."


def _ranking_observation_text(rankings: List[Dict[str, Any]]) -> str:
    if not rankings:
        return "The ranking tables for this graph were limited, so the strongest and weakest cases were inferred from the available curve summaries only."
    row = rankings[0]
    ordered = row.get("ordered_best_to_worst", []) or []
    if not ordered:
        return "The ranking tables for this graph were limited, so the strongest and weakest cases were inferred from the available curve summaries only."
    best = ordered[0]
    worst = ordered[-1]
    middle = ", ".join(str(item.get("router", "")) for item in ordered[1:4] if item.get("router"))
    parts = [
        f"Best-ranked case in the supplied ranking table: {best.get('router', '—')} (rank {best.get('rank', '—')}, avg {best.get('avg_actual', '—')}).",
        f"Lowest-ranked case in the supplied ranking table: {worst.get('router', '—')} (rank {worst.get('rank', '—')}, avg {worst.get('avg_actual', '—')}).",
    ]
    if middle:
        parts.append(f"Other prominent ranked routers in the same view: {middle}.")
    return " ".join(parts)


def _fallback_metric_patterns(metric_context: Dict[str, Any]) -> str:
    curves = metric_context.get("curve_summaries", []) or []
    if not curves:
        return "The metric-level pattern summary is limited because the curve summaries were not available."
    top = curves[:3]
    bullets = []
    for row in top:
        bullets.append(
            f"- {row.get('router_key', '—')} | {row.get('floor_name', '—')} | {row.get('band', '—')}: avg {row.get('avg_curve_value', '—')}, trend {row.get('trend', '—')}."
        )
    return "\n".join(bullets)


def build_graph_prompt(
    router_name: str,
    metric_display: str,
    metric_context: Dict[str, Any],
    plot_info: Dict[str, Any],
    extra_instructions: str = "",
) -> str:
    focus_curves = _filter_relevant_curves(metric_context, plot_info)
    focus_rankings = _filter_relevant_rankings(metric_context, plot_info)
    preference = "higher values are better" if metric_context.get("higher_is_better", True) else "lower values are better"
    payload = {
        "router_name": router_name,
        "metric_display": metric_display,
        "metric_preference": preference,
        "plot_info": plot_info,
        "relevant_curve_summaries": focus_curves,
        "relevant_rankings": focus_rankings,
        "dominant_distance_behavior": _distance_behavior_from_curves(focus_curves, metric_context.get("higher_is_better", True)),
    }
    if extra_instructions.strip():
        payload["extra_instructions"] = extra_instructions.strip()

    instructions = (
        "Write a formal report for a single comparative-study graph.\n"
        "Non-negotiable rules:\n"
        "- Use only the supplied numeric context.\n"
        "- Never say the metric improves with distance unless the supplied trend data explicitly supports that.\n"
        "- If most curves decline or worsen with distance, state that clearly.\n"
        "- Do not add questions, offers to refine, chatty endings, or troubleshooting suggestions.\n"
        "- Keep each section concise and factual.\n"
    )
    output_format = (
        "Format exactly with these headings:\n"
        "# <Graph title>\n"
        "## What the graph compares\n"
        "## Main observations\n"
        "## Distance behavior\n"
        "## Engineering reasoning\n"
        "## Cautions\n"
    )
    return _truncate_text(
        f"{instructions}\n{output_format}\nJSON context:\n{json.dumps(payload, indent=2, ensure_ascii=False)}"
    )


def build_metric_overview_prompt(
    router_name: str,
    metric_display: str,
    metric_context: Dict[str, Any],
    graph_report_excerpts: List[Dict[str, Any]],
    extra_instructions: str = "",
) -> str:
    preference = "higher values are better" if metric_context.get("higher_is_better", True) else "lower values are better"
    payload = {
        "router_name": router_name,
        "metric_display": metric_display,
        "metric_preference": preference,
        "metric_context": metric_context,
        "graph_reports": graph_report_excerpts,
        "distance_behavior_summary": _distance_behavior_from_curves(metric_context.get("curve_summaries", []) or [], metric_context.get("higher_is_better", True)),
    }
    if extra_instructions.strip():
        payload["extra_instructions"] = extra_instructions.strip()
    return _truncate_text(
        "Write one metric-level report across all graph summaries for this metric. "
        "Use only the supplied context. Do not add questions or assistant-style endings.\n\n"
        "Format exactly with these headings:\n"
        "# <Metric> Overall Summary\n"
        "## Strong patterns\n"
        "## Weak patterns\n"
        "## Distance-related findings\n"
        "## Engineering takeaways\n"
        "## Data limits\n\n"
        f"JSON context:\n{json.dumps(payload, indent=2, ensure_ascii=False)}",
        limit=22000,
    )


def build_router_overall_prompt(
    router_name: str,
    metric_summaries: List[Dict[str, Any]],
    extra_instructions: str = "",
) -> str:
    payload = {"router_name": router_name, "metric_summaries": metric_summaries}
    if extra_instructions.strip():
        payload["extra_instructions"] = extra_instructions.strip()
    return _truncate_text(
        "Write one consolidated report across all selected metrics, floors, routers, and bands. "
        "Stay grounded in the supplied metric summaries only. "
        "Do not add questions, troubleshooting invitations, or assistant-style wrap-up text.\n\n"
        "Format exactly with these headings:\n"
        "# <Router> Overall Comparative Study Summary\n"
        "## Scope covered\n"
        "## Cross-metric findings\n"
        "## Floors and bands\n"
        "## Router comparisons\n"
        "## Engineering takeaways\n"
        "## Data limits\n\n"
        f"JSON context:\n{json.dumps(payload, indent=2, ensure_ascii=False)}",
        limit=24000,
    )


def _normalize_graph_report(
    raw_text: str,
    router_name: str,
    metric_display: str,
    metric_context: Dict[str, Any],
    plot_info: Dict[str, Any],
) -> str:
    cleaned = _clean_generated_report_text(raw_text)
    _, sections = _extract_markdown_sections(cleaned)
    focus_curves = _filter_relevant_curves(metric_context, plot_info)
    focus_rankings = _filter_relevant_rankings(metric_context, plot_info)
    title = f"{metric_display} Graph Summary - {Path(str(plot_info.get('plot_name', 'graph'))).stem}"

    normalized = {
        "What the graph compares": sections.get("What the graph compares") or _graph_scope_text(plot_info),
        "Main observations": sections.get("Main observations") or _ranking_observation_text(focus_rankings),
        "Distance behavior": _distance_behavior_from_curves(focus_curves, metric_context.get("higher_is_better", True)),
        "Engineering reasoning": sections.get("Engineering reasoning") or (
            "The engineering interpretation should follow the supplied rankings and curve trends only. "
            "Where curves degrade with distance, the usable performance envelope is tightening rather than improving."
        ),
        "Cautions": sections.get("Cautions") or (
            "This report is limited to the available curve summaries, ranking tables, and generated graph context. "
            "It does not infer conditions that are not present in the provided data."
        ),
    }
    return _compose_markdown_report(title, GRAPH_SECTION_ORDER, normalized)


def _normalize_metric_overview(
    raw_text: str,
    router_name: str,
    metric_display: str,
    metric_context: Dict[str, Any],
) -> str:
    cleaned = _clean_generated_report_text(raw_text)
    _, sections = _extract_markdown_sections(cleaned)
    title = f"{metric_display} Overall Summary"
    normalized = {
        "Strong patterns": sections.get("Strong patterns") or _fallback_metric_patterns(metric_context),
        "Weak patterns": sections.get("Weak patterns") or (
            "The weakest patterns were those with poorer average curve values or less stable distance behavior in the available summaries."
        ),
        "Distance-related findings": _distance_behavior_from_curves(metric_context.get("curve_summaries", []) or [], metric_context.get("higher_is_better", True)),
        "Engineering takeaways": sections.get("Engineering takeaways") or (
            "Engineering conclusions should be based on the floors, routers, and bands that remain consistently strong across distance rather than on isolated points."
        ),
        "Data limits": sections.get("Data limits") or (
            "This overview is constrained by the generated plots, ranking tables, and curve summaries available for this metric."
        ),
    }
    return _compose_markdown_report(title, METRIC_SECTION_ORDER, normalized)


def _normalize_router_overall_report(raw_text: str, router_name: str, metric_summaries: List[Dict[str, Any]]) -> str:
    cleaned = _clean_generated_report_text(raw_text)
    _, sections = _extract_markdown_sections(cleaned)
    covered_metrics = ", ".join(str(item.get("metric_display", "—")) for item in metric_summaries if item.get("metric_display"))
    floors = sorted({str(f) for item in metric_summaries for f in item.get("floors", []) if f})
    bands = sorted({str(b) for item in metric_summaries for b in item.get("bands", []) if b})
    scope_text = (
        f"Metrics covered: {covered_metrics or 'none identified'}. "
        f"Floors covered: {', '.join(floors) or 'not identified'}. "
        f"Bands covered: {', '.join(bands) or 'not identified'}."
    )
    router_compare_lines = []
    for item in metric_summaries:
        ranking_excerpt = item.get("ranking_excerpt", "")
        if ranking_excerpt:
            router_compare_lines.append(f"- {item.get('metric_display', 'Metric')}: {ranking_excerpt}")
    normalized = {
        "Scope covered": sections.get("Scope covered") or scope_text,
        "Cross-metric findings": sections.get("Cross-metric findings") or (
            "Cross-metric interpretation should prioritize recurring patterns that appear across multiple floors, routers, and bands rather than isolated wins."
        ),
        "Floors and bands": sections.get("Floors and bands") or scope_text,
        "Router comparisons": sections.get("Router comparisons") or ("\n".join(router_compare_lines) if router_compare_lines else "Router comparison statements were limited to the ranking evidence available inside the metric summaries."),
        "Engineering takeaways": sections.get("Engineering takeaways") or (
            "The strongest comparative result is the one that remains stable across distance, floors, and bands rather than only in one narrow case."
        ),
        "Data limits": sections.get("Data limits") or (
            "This consolidated report is limited to the selected metric outputs and their generated summaries."
        ),
    }
    title = f"{router_name} Overall Comparative Study Summary"
    return _compose_markdown_report(title, OVERALL_SECTION_ORDER, normalized)


def generate_graph_ai_report(
    router_name: str,
    metric_display: str,
    metric_context: Dict[str, Any],
    plot_path: Path,
    ai_output_root: Path,
    cfg: LocalAIConfig,
) -> Dict[str, Any]:
    plot_info = infer_plot_info(plot_path)
    prompt = build_graph_prompt(router_name, metric_display, metric_context, plot_info, extra_instructions=cfg.extra_instructions)
    raw_text = run_local_ai(prompt, cfg)
    text = _normalize_graph_report(raw_text, router_name, metric_display, metric_context, plot_info)

    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", plot_path.stem.lower()).strip("_") + "__graph_summary"
    files = save_text_artifacts(
        text=text,
        title=f"{metric_display} Graph Summary - {plot_path.stem}",
        out_dir=ai_output_root,
        stem=stem,
        metadata={
            "router_name": router_name,
            "metric_display": metric_display,
            "plot_path": str(plot_path),
            "plot_info": plot_info,
            "config": asdict(cfg),
            "prompt": prompt,
            "raw_text": raw_text,
        },
        image_paths=[plot_path],
    )

    return {
        "plot_name": plot_path.name,
        "plot_path": str(plot_path),
        "plot_info": plot_info,
        "prompt": prompt,
        "raw_text": raw_text,
        "text": text,
        "files": files,
    }


def generate_metric_graph_reports(
    router_name: str,
    metric_display: str,
    metric_output_dir: Path,
    ai_output_root: Path,
    cfg: LocalAIConfig,
) -> Dict[str, Any]:
    metric_context = build_metric_context(metric_output_dir, metric_display)
    plot_paths = sorted(metric_output_dir.rglob("*.png"))
    graph_reports: List[Dict[str, Any]] = []

    for plot_path in plot_paths:
        graph_reports.append(
            generate_graph_ai_report(
                router_name=router_name,
                metric_display=metric_display,
                metric_context=metric_context,
                plot_path=plot_path,
                ai_output_root=ai_output_root / "graphs",
                cfg=cfg,
            )
        )

    overview = None
    if graph_reports:
        excerpts = [
            {
                "plot_name": item.get("plot_name"),
                "plot_info": item.get("plot_info"),
                "report_excerpt": _truncate_text(str(item.get("text", "")), limit=2200),
            }
            for item in graph_reports
        ]
        prompt = build_metric_overview_prompt(router_name, metric_display, metric_context, excerpts, extra_instructions=cfg.extra_instructions)
        raw_text = run_local_ai(prompt, cfg)
        text = _normalize_metric_overview(raw_text, router_name, metric_display, metric_context)
        files = save_text_artifacts(
            text=text,
            title=f"{metric_display} Overall Summary",
            out_dir=ai_output_root,
            stem=re.sub(r"[^A-Za-z0-9._-]+", "_", metric_display.lower()) + "__overall_summary",
            metadata={
                "router_name": router_name,
                "metric_display": metric_display,
                "metric_output_dir": str(metric_output_dir),
                "metric_context": metric_context,
                "graph_report_count": len(graph_reports),
                "prompt": prompt,
                "config": asdict(cfg),
                "raw_text": raw_text,
            },
            image_paths=plot_paths,
        )
        overview = {"prompt": prompt, "raw_text": raw_text, "text": text, "files": files}

    floors = sorted({str(row.get("floor_name", "")) for row in metric_context.get("curve_summaries", []) if row.get("floor_name")})
    bands = sorted({str(row.get("band", "")) for row in metric_context.get("curve_summaries", []) if row.get("band")})
    ranking_excerpt = ""
    rankings = metric_context.get("rankings", []) or []
    if rankings and rankings[0].get("ordered_best_to_worst"):
        best = rankings[0]["ordered_best_to_worst"][0]
        worst = rankings[0]["ordered_best_to_worst"][-1]
        ranking_excerpt = f"best observed rank {best.get('router', '—')}; lowest observed rank {worst.get('router', '—')}"

    return {
        "metric_display": metric_display,
        "metric_output_dir": str(metric_output_dir),
        "metric_context": metric_context,
        "graph_reports": graph_reports,
        "overview": overview,
        "plot_paths": [str(p) for p in plot_paths],
        "floors": floors,
        "bands": bands,
        "ranking_excerpt": ranking_excerpt,
    }


def generate_router_overall_report(
    router_name: str,
    metric_results: List[Dict[str, Any]],
    ai_output_root: Path,
    cfg: LocalAIConfig,
) -> Optional[Dict[str, Any]]:
    ok_results = [item for item in metric_results if item]
    if not ok_results:
        return None

    metric_summaries: List[Dict[str, Any]] = []
    image_paths: List[Path] = []
    for item in ok_results:
        overview_text = ""
        overview = item.get("overview")
        if isinstance(overview, dict):
            overview_text = str(overview.get("text", ""))
        metric_summaries.append(
            {
                "metric_display": item.get("metric_display"),
                "floors": item.get("floors", []),
                "bands": item.get("bands", []),
                "curve_summary_count": len(item.get("metric_context", {}).get("curve_summaries", []) or []),
                "ranking_excerpt": item.get("ranking_excerpt", ""),
                "overview_excerpt": _truncate_text(overview_text, limit=2600),
            }
        )
        image_paths.extend([Path(p) for p in item.get("plot_paths", []) if Path(p).exists()])

    prompt = build_router_overall_prompt(router_name, metric_summaries, extra_instructions=cfg.extra_instructions)
    raw_text = run_local_ai(prompt, cfg)
    text = _normalize_router_overall_report(raw_text, router_name, metric_summaries)
    files = save_text_artifacts(
        text=text,
        title=f"{router_name} Overall Comparative Study Summary",
        out_dir=ai_output_root,
        stem="router_overall_summary",
        metadata={
            "router_name": router_name,
            "metric_summaries": metric_summaries,
            "prompt": prompt,
            "config": asdict(cfg),
            "raw_text": raw_text,
        },
        image_paths=image_paths,
    )
    return {"prompt": prompt, "raw_text": raw_text, "text": text, "files": files, "image_count": len(image_paths)}


# -----------------------------
# Knowledge-base chat
# -----------------------------
def _tokenize(text: str) -> List[str]:
    return re.findall(r"[A-Za-z0-9_.-]+", (text or "").lower())


def _chunk_text(text: str, chunk_size: int = 1400, overlap: int = 180) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]
    out = []
    start = 0
    while start < len(text):
        out.append(text[start : start + chunk_size])
        start += max(1, chunk_size - overlap)
    return out


def _csv_to_text(path: Path, max_rows: int = 24) -> str:
    try:
        df = pd.read_csv(path)
    except Exception:
        return ""
    if df.empty:
        return f"CSV file {path.name} is empty."
    head = df.head(max_rows)
    return f"CSV file: {path.name}\nColumns: {', '.join(map(str, head.columns))}\nSample rows:\n{head.to_csv(index=False)}"


def build_knowledge_base(router_dir: Path) -> List[Dict[str, Any]]:
    roots = [router_dir / "ai_reports", router_dir / "rvr_outputs", router_dir / "csv_outputs"]
    docs: List[Dict[str, Any]] = []
    for root in roots:
        if not root.exists():
            continue
        for path in sorted(root.rglob("*")):
            if not path.is_file():
                continue
            suffix = path.suffix.lower()
            text = ""
            if suffix in {".md", ".txt", ".log"}:
                try:
                    text = path.read_text(encoding="utf-8", errors="ignore")
                except Exception:
                    text = ""
            elif suffix == ".json":
                try:
                    raw = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
                    text = json.dumps(raw, indent=2, ensure_ascii=False)
                except Exception:
                    text = ""
            elif suffix == ".csv":
                text = _csv_to_text(path)
            else:
                continue
            if not text.strip():
                continue
            rel = str(path.relative_to(router_dir))
            for idx, chunk in enumerate(_chunk_text(text)):
                docs.append({"source": rel, "chunk_id": idx, "text": chunk, "tokens": set(_tokenize(chunk))})
    return docs


def retrieve_kb_chunks(query: str, kb_docs: List[Dict[str, Any]], top_k: int = 8) -> List[Dict[str, Any]]:
    q_tokens = set(_tokenize(query))
    scored: List[Tuple[float, Dict[str, Any]]] = []
    for doc in kb_docs:
        d_tokens = doc.get("tokens", set())
        if not d_tokens:
            continue
        overlap = len(q_tokens & d_tokens)
        if overlap == 0:
            continue
        score = overlap / (1.0 + math.log(2 + len(d_tokens)))
        if any(tok in doc.get("source", "").lower() for tok in q_tokens):
            score += 0.5
        scored.append((score, doc))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [item[1] for item in scored[:top_k]]


def build_kb_chat_prompt(
    router_name: str,
    question: str,
    retrieved_chunks: List[Dict[str, Any]],
    history: Optional[List[Dict[str, str]]] = None,
    extra_instructions: str = "",
) -> str:
    history = history or []
    trimmed_history = history[-6:]
    context_blobs = []
    for item in retrieved_chunks:
        context_blobs.append(f"[Source: {item.get('source')} | Chunk: {item.get('chunk_id')}]\n{item.get('text')}")
    payload = {"router_name": router_name, "question": question, "history": trimmed_history, "sources": context_blobs}
    if extra_instructions.strip():
        payload["extra_instructions"] = extra_instructions.strip()
    return _truncate_text(
        "Answer the user's question using only the supplied knowledge-base context from the WiFi site survey project. "
        "Be direct and practical. If the answer is not supported by the context, say that clearly. "
        "At the end, include a short 'Sources used' line listing the source file names you relied on.\n\n"
        f"JSON context:\n{json.dumps(payload, indent=2, ensure_ascii=False)}",
        limit=22000,
    )


def chat_with_knowledge_base(
    router_name: str,
    router_dir: Path,
    question: str,
    cfg: LocalAIConfig,
    history: Optional[List[Dict[str, str]]] = None,
    kb_docs: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    docs = kb_docs if kb_docs is not None else build_knowledge_base(router_dir)
    retrieved = retrieve_kb_chunks(question, docs, top_k=8)
    prompt = build_kb_chat_prompt(router_name, question, retrieved, history=history, extra_instructions=cfg.extra_instructions)
    answer = run_local_ai(prompt, cfg)
    return {
        "answer": answer,
        "sources": [item.get("source") for item in retrieved],
        "retrieved": retrieved,
        "prompt": prompt,
        "kb_doc_count": len(docs),
    }
